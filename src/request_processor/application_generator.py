"""
Генерация заявки на проведение испытаний по шаблону Word (.docx).
"""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document

from .sqlite_repo import (
    DB_PATH_DEFAULT,
    GENERATED_DIR_DEFAULT,
    PROJECT_ROOT,
    get_order_details,
)

TEMPLATE_DEFAULT = PROJECT_ROOT / "data" / "templates" / "zayavka_ispytaniy.docx"
TEST_TYPE_DEFAULT = "периодические"


def _unique_cells(row) -> list:
    seen: list = []
    for cell in row.cells:
        if cell._tc not in [c._tc for c in seen]:
            seen.append(cell)
    return seen


def _set_row_value(table, row_idx: int, value: str, *, col: int = 1) -> None:
    if row_idx >= len(table.rows):
        return
    cells = _unique_cells(table.rows[row_idx])
    if len(cells) > col:
        cells[col].text = value or "—"


def _format_criteria(lines: list[dict]) -> str:
    if not lines:
        return "Согласно объёму испытаний, указанному в приложении к заявке."
    parts = []
    for line in lines:
        name = (line.get("test_name") or "").strip()
        if name:
            parts.append(name)
    if not parts:
        return "Согласно объёму испытаний, указанному в приложении к заявке."
    return "; ".join(parts)


def _detect_test_type(subject: str | None) -> str:
    if not subject:
        return TEST_TYPE_DEFAULT
    s = subject.lower()
    if "сертификац" in s:
        return "сертификационные"
    if "контрольн" in s:
        return "контрольные"
    if "периодич" in s:
        return "периодические"
    return TEST_TYPE_DEFAULT


def _append_table_row(table) -> None:
    """Добавляет строку в таблицу (клон последней)."""
    tbl = table._tbl
    new_tr = deepcopy(tbl.tr_lst[-1])
    tbl.append(new_tr)


def _fill_appendix_table(
    table,
    marks: list[dict],
    *,
    get_lines,
    get_document,
) -> None:
    """Заполняет приложение: марки, ТУ, критерии (виды испытаний)."""
    while len(table.rows) > 1:
        table._tbl.remove(table._tbl.tr_lst[-1])

    for idx, mark_row in enumerate(marks, start=1):
        calc_id = mark_row.get("calculation_id")
        mark_name = mark_row.get("mark") or ""
        document = get_document(mark_name) or "—"
        lines = get_lines(int(calc_id)) if calc_id else []
        criteria = _format_criteria(lines)

        _append_table_row(table)

        row = table.rows[-1]
        ucells = _unique_cells(row)
        if len(ucells) >= 4:
            ucells[0].text = str(idx)
            ucells[1].text = mark_name
            ucells[2].text = document
            ucells[3].text = criteria


def generate_application_from_order(
    order_id: int,
    output_path: Path | str | None = None,
    *,
    template_path: Path | str = TEMPLATE_DEFAULT,
    db_path: Path | str = DB_PATH_DEFAULT,
) -> Path:
    """
    Формирует заявку на испытания по заказу: лист 1 — форма, лист 2 — объём испытаний.
    """
    from .sqlite_repo import (
        get_calculation_lines,
        get_cable_mark_document,
        get_organization_by_id,
        update_order_application_path,
    )

    details = get_order_details(order_id, db_path=db_path)
    if not details:
        raise ValueError(f"Заказ №{order_id} не найден")

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Шаблон заявки не найден: {template}")

    if output_path is None:
        customer = details.get("customer_name") or "заказчик"
        safe = re.sub(r'[<>:"/\\|?*«»]', "_", customer).strip("._ ")[:30] or "заказчик"
        out_dir = GENERATED_DIR_DEFAULT
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / f"Заявка_{safe}_заказ{order_id}_{datetime.now():%Y%m%d_%H%M}.docx"
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    shutil.copy2(template, output_path)
    doc = Document(str(output_path))

    if len(doc.tables) < 2:
        raise ValueError("Шаблон должен содержать 2 таблицы (форма + приложение)")

    form = doc.tables[0]
    appendix = doc.tables[1]

    order_no = str(order_id)
    _set_row_value(form, 0, f"ЗАЯВКА № {order_no}", col=0)

    customer_name = details.get("customer_name") or "—"
    customer_ral = "—"
    if details.get("customer_org_id"):
        org = get_organization_by_id(int(details["customer_org_id"]), db_path)
        if org and org.get("fsa_registry_number"):
            customer_ral = org["fsa_registry_number"]
        elif org and org.get("is_accredited"):
            customer_ral = "аккредитован"

    _set_row_value(form, 4, customer_name)
    _set_row_value(form, 5, customer_ral)
    _set_row_value(form, 6, details.get("customer_address") or "—")
    _set_row_value(form, 7, details.get("customer_address") or "—")
    if details.get("customer_org_id"):
        org = get_organization_by_id(int(details["customer_org_id"]), db_path)
        if org:
            _set_row_value(form, 8, org.get("phone") or "—")
            _set_row_value(form, 9, org.get("email") or "—")

    mfg_name = details.get("manufacturer_name") or customer_name
    _set_row_value(form, 11, mfg_name)
    mfg_addr = details.get("manufacturer_address") or details.get("customer_address") or "—"
    _set_row_value(form, 12, mfg_addr)
    _set_row_value(form, 13, mfg_addr)
    if details.get("manufacturer_org_id"):
        morg = get_organization_by_id(int(details["manufacturer_org_id"]), db_path)
        if morg:
            _set_row_value(form, 14, morg.get("phone") or "—")
            _set_row_value(form, 15, morg.get("email") or "—")

    test_type = _detect_test_type(details.get("subject"))
    _set_row_value(form, 17, test_type)

    marks = details.get("marks") or []
    docs = sorted(
        {get_cable_mark_document(m.get("mark") or "", db_path) or "" for m in marks}
        - {""}
    )
    if docs:
        _set_row_value(form, 22, ", ".join(docs))

    def _lines_fn(calc_id: int) -> list[dict]:
        return get_calculation_lines(calc_id, db_path=db_path)

    def _doc_fn(mark: str) -> str | None:
        return get_cable_mark_document(mark, db_path=db_path)

    _fill_appendix_table(appendix, marks, get_lines=_lines_fn, get_document=_doc_fn)

    doc.save(str(output_path))
    update_order_application_path(order_id, str(output_path.resolve()), db_path=db_path)
    return Path(output_path).resolve()