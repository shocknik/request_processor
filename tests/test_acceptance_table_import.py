"""Волна 2: парсер таблиц приёмки + импорт эталонов."""

from __future__ import annotations

from pathlib import Path

import pytest
from click.testing import CliRunner
from docx import Document

from request_processor.cli import cli
from request_processor.generation.acceptance_table_import import (
    expand_clause_refs,
    import_acceptance_docx,
    import_etalon_batch,
    parse_acceptance_from_docx,
    parse_acceptance_from_raw_text_methods,
    RAW_TEXT_DIR,
    TU_CORPUS,
)
from request_processor.persistence.sqlite_repo import (
    init_db,
    list_acceptance_items,
    show_norm_catalog,
)


def test_expand_clause_range() -> None:
    assert expand_clause_refs("2.2.2, 2.3.1 - 2.3.6") == [
        "2.2.2",
        "2.3.1",
        "2.3.2",
        "2.3.3",
        "2.3.4",
        "2.3.5",
        "2.3.6",
    ]
    assert expand_clause_refs("2.5.1") == ["2.5.1"]
    assert "2.10.5" in expand_clause_refs("2.10.1 – 2.10.5")


def _mini_acceptance_docx(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Таблица 4 – Состав приемо-сдаточных испытаний")
    t = doc.add_table(rows=4, cols=4)
    hdr = (
        ("Группа испытаний", "Вид испытания или проверки", "технических требований", "методов контроля"),
        ("С1", "Измерение коэффициента затухания", "2.4", "5.3"),
        ("С1", "Проверка конструкции", "2.3.1 - 2.3.3", "5.2"),
        ("С2", "Проверка маркировки и упаковки", "2.8.1, 2.9", "5.7.1"),
    )
    for i, row in enumerate(hdr):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    doc.add_paragraph("Таблица 5 – Состав периодических испытаний")
    t2 = doc.add_table(rows=2, cols=4)
    t2.rows[0].cells[0].text = "Группа"
    t2.rows[0].cells[1].text = "Вид испытания или проверки"
    t2.rows[0].cells[2].text = "технических требований"
    t2.rows[0].cells[3].text = "методов контроля"
    t2.rows[1].cells[0].text = "П1"
    t2.rows[1].cells[1].text = "Прочность к растягивающему усилию"
    t2.rows[1].cells[2].text = "2.5.1"
    t2.rows[1].cells[3].text = "5.4.1"
    doc.save(path)
    return path


def test_parse_mini_docx(tmp_path: Path) -> None:
    fp = _mini_acceptance_docx(tmp_path / "tu_test.docx")
    cat = parse_acceptance_from_docx(fp)
    assert len(cat.items) >= 3
    names = {i.name_exact for i in cat.items}
    assert any("затухан" in n.lower() for n in names)
    stretch = next(i for i in cat.items if "растягивающ" in i.name_exact.lower())
    assert stretch.requirement_clauses == ["2.5.1"]
    assert stretch.method_clauses == ["5.4.1"]
    assert stretch.test_category == "periodic"
    mark = next(i for i in cat.items if "маркиров" in i.name_exact.lower())
    assert not mark.billable
    # диапазон развернут
    constr = next(i for i in cat.items if "конструкц" in i.name_exact.lower())
    assert "2.3.1" in constr.requirement_clauses
    assert "2.3.3" in constr.requirement_clauses


def test_import_mini_docx_to_db(tmp_path: Path) -> None:
    db = tmp_path / "a.db"
    init_db(db)
    fp = _mini_acceptance_docx(tmp_path / "tu2.docx")
    r = import_acceptance_docx(
        fp, db_path=db, replace=True, doc_id="ТУ-TEST-MINI"
    )
    assert r["items"] >= 3
    items = list_acceptance_items(doc_id="ТУ-TEST-MINI", db_path=db)
    assert len(items) >= 3
    cat = show_norm_catalog(doc_id="ТУ-TEST-MINI", db_path=db)
    assert cat is not None
    assert len(cat["acceptance_items"]) >= 3


def test_cli_import_docx(tmp_path: Path) -> None:
    db = tmp_path / "c.db"
    init_db(db)
    fp = _mini_acceptance_docx(tmp_path / "tu3.docx")
    runner = CliRunner()
    res = runner.invoke(
        cli,
        [
            "import-acceptance-docx",
            "--file",
            str(fp),
            "--doc-id",
            "ТУ-CLI-1",
            "--db",
            str(db),
            "--no-match-price",
        ],
    )
    assert res.exit_code == 0, res.output
    assert "items=" in res.output


@pytest.mark.skipif(
    not list(TU_CORPUS.glob("*131*2025*.docx")),
    reason="локальный корпус ТУ 131 не установлен",
)
def test_real_131_parse_smoke() -> None:
    path = list(TU_CORPUS.glob("*131*2025*.docx"))[0]
    cat = parse_acceptance_from_docx(path)
    assert len(cat.items) >= 10
    assert any("растягивающ" in i.name_exact.lower() for i in cat.items)
    assert not any(i.name_exact.upper().startswith("ГОСТ") for i in cat.items)


@pytest.mark.skipif(
    not list(RAW_TEXT_DIR.glob("*005-01*.txt")),
    reason="raw_text 005 отсутствует",
)
def test_real_005_raw_text_smoke() -> None:
    path = list(RAW_TEXT_DIR.glob("*005-01*.txt"))[0]
    cat = parse_acceptance_from_raw_text_methods(path, doc_id="ТУ 16.К99-005-01")
    assert len(cat.items) >= 5
    assert any("затухан" in i.name_exact.lower() for i in cat.items)


@pytest.mark.skipif(
    not list(TU_CORPUS.glob("*131*2025*.docx")),
    reason="локальный корпус не установлен",
)
def test_etalon_batch_import(tmp_path: Path) -> None:
    db = tmp_path / "etalon.db"
    init_db(db)
    results = import_etalon_batch(db_path=db, replace=True)
    ok = [r for r in results if not r.get("error") and r.get("items", 0) > 0]
    assert len(ok) >= 2  # как минимум 131 и 141
    total = sum(int(r.get("items") or 0) for r in ok)
    assert total >= 20
