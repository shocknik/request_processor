"""Импорт программ испытаний из DOCX — марки, таблицы, списки."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from request_processor.generation.program_importer import (
    import_program_from_docx,
    parse_program_docx,
)
from request_processor.persistence.sqlite_repo import (
    get_test_program,
    init_db,
    list_test_programs,
    match_program_items_to_price,
)


def _make_single_mark_program(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("УТВЕРЖДАЮ")
    doc.add_paragraph("ПРОГРАММА")
    doc.add_paragraph(
        "исследовательских испытаний кабеля симметричного парной скрутки "
        "марки СПЕЦЛАН F/UTP Cat 5е PUR 4x2x0,52, изготовленного и представленного "
        "на испытания по ТУ 16.К99-058-2014"
    )
    doc.add_paragraph("1. Объект испытаний.")
    doc.add_paragraph(
        "В качестве типового представителя выбрана марка "
        "СПЕЦЛАН F/UTP Cat 5е PUR 4x2x0,52 – кабель связи, категории 5е."
    )
    table = doc.add_table(rows=5, cols=4)
    rows = [
        ("№ п/п", "Вид испытаний и проверок", "Пункты ТУ 16.К99-058-2014", ""),
        ("№ п/п", "Вид испытаний и проверок", "технических требований", "методов испытаний"),
        ("1.", "Проверка конструкции и конструктивных размеров", "1.2.2, 1.2.3", "4.2.1"),
        (
            "2.",
            "Проверка герметичности изоляции Проверка сплошности экрана",
            "1.3.3 1.3.11",
            "4.2.2 4.2.3",
        ),
        ("3.", "Определение электрического сопротивления жилы", "1.4.1", "4.3.1"),
    ]
    for ri, vals in enumerate(rows):
        for ci, v in enumerate(vals):
            table.rows[ri].cells[ci].text = v
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _make_multi_mark_list_program(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("ПРОГРАММА")
    doc.add_paragraph(
        "приемочных испытаний кабелей универсальных, марок"
    )
    doc.add_paragraph("1. Объект испытаний.")
    doc.add_paragraph("В качестве типового представителя выбраны марки кабеля:")
    doc.add_paragraph(
        "1) СКАБ-СЭ 660Кнг(А)-LS-ХЛ 61х0,35 л4 – кабель для экстремальных условий"
    )
    doc.add_paragraph(
        "2) СКАБ-СЭом 660БсПВнг(А)-LS 1х16 м5 – кабель одножильный"
    )
    doc.add_paragraph(
        "Образцы по ТУ 27.32.13-099-47273194-2020 «Кабели универсальные»"
    )
    table = doc.add_table(rows=3, cols=5)
    rows = [
        (
            "№ п/п",
            "Вид испытаний и проверок",
            "технических требований",
            "методов испытаний",
            "Применяемость",
        ),
        ("1.", "Проверка конструкции", "1.2.2", "4.2.1", "1, 2"),
        ("2.", "Испытание напряжением", "1.4.1", "4.3.2", "1, 2"),
    ]
    for ri, vals in enumerate(rows):
        for ci, v in enumerate(vals):
            table.rows[ri].cells[ci].text = v
    doc.save(str(path))
    return path


def test_parse_single_mark_and_split_jammed(tmp_path: Path) -> None:
    docx = _make_single_mark_program(tmp_path / "one.docx")
    parsed = parse_program_docx(docx)
    assert "СПЕЦЛАН" in parsed.cable_mark_text.upper() or any(
        "СПЕЦЛАН" in m.upper() for m in parsed.cable_marks
    )
    assert "16.К99" in (parsed.tu_ref or "")
    # at least 3 logical tests; jammed row may split to 2
    assert len(parsed.items) >= 3
    names = " ".join(it.name for it in parsed.items)
    assert "герметич" in names.lower() or "сопротивления" in names.lower()


def test_parse_multi_mark_list(tmp_path: Path) -> None:
    docx = _make_multi_mark_list_program(tmp_path / "multi.docx")
    parsed = parse_program_docx(docx)
    assert len(parsed.cable_marks) >= 2
    assert any("СКАБ-СЭ" in m for m in parsed.cable_marks)
    assert any("СКАБ-СЭом" in m or "СКАБ" in m for m in parsed.cable_marks)
    assert len(parsed.items) >= 2
    assert "27.32.13" in (parsed.tu_ref or "")


def test_import_program_to_db(tmp_path: Path) -> None:
    db = tmp_path / "p.db"
    init_db(db)
    docx = _make_single_mark_program(tmp_path / "prog.docx")
    result = import_program_from_docx(docx, db_path=db, match_price=True)
    assert result["program_id"] >= 1
    assert result["items_count"] >= 2
    assert result.get("cable_marks") or result.get("cable_mark_text")
    prog = get_test_program(result["program_id"], db_path=db)
    assert prog is not None
    assert len(prog["items"]) >= 2


def test_match_price_runs(tmp_path: Path) -> None:
    db = tmp_path / "p.db"
    init_db(db)
    docx = _make_single_mark_program(tmp_path / "prog.docx")
    result = import_program_from_docx(docx, db_path=db, match_price=False)
    stats = match_program_items_to_price(result["program_id"], db_path=db)
    assert "matched" in stats and "unmatched" in stats
