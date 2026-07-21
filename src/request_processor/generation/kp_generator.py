"""
Генерация коммерческого предложения (КП) в формате Word.
Стили: classic | modern | compact. Логотип и реквизиты — lab_profile.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..models import CommercialProposal, KPMarkLine
from .lab_profile import LabProfile, load_lab_profile

LAB_NAME = "ООО «Испытательный центр»"
LAB_TAGLINE = "Испытательный центр кабельной продукции"
DEFAULT_VALIDITY_DAYS = 30
VAT_PERCENT = 22

KpStyle = Literal["classic", "modern", "compact"]

_STYLE_COLORS = {
    "classic": {
        "title": RGBColor(0x25, 0x63, 0xEB),
        "header_bg": "2563EB",
        "total_bg": "E2E8F0",
        "name": RGBColor(0x1E, 0x3A, 0x5F),
        "muted": RGBColor(0x64, 0x74, 0x8B),
    },
    "modern": {
        "title": RGBColor(0x0F, 0x76, 0x6E),
        "header_bg": "0F766E",
        "total_bg": "CCFBF1",
        "name": RGBColor(0x13, 0x40, 0x3C),
        "muted": RGBColor(0x5E, 0x6B, 0x73),
    },
    "compact": {
        "title": RGBColor(0x33, 0x41, 0x55),
        "header_bg": "334155",
        "total_bg": "F1F5F9",
        "name": RGBColor(0x1E, 0x29, 0x3B),
        "muted": RGBColor(0x64, 0x74, 0x8B),
    },
}


def format_money(amount: float) -> str:
    """1 234 567,89"""
    whole, frac = f"{amount:.2f}".split(".")
    grouped = ""
    for i, ch in enumerate(reversed(whole)):
        if i and i % 3 == 0:
            grouped = " " + grouped
        grouped = ch + grouped
    return f"{grouped},{frac}"


def build_intro_text(subject: str, customer: str) -> str:
    subject = subject.strip() or "Проведение испытаний"
    customer = customer.strip()
    if customer:
        return f"{subject} для {customer}"
    return subject


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, *, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_header_block(
    doc: Document,
    profile: LabProfile,
    colors: dict,
    *,
    style: str,
) -> None:
    """Шапка КП: логотип (КТ) + реквизиты; solid layout для всех стилей."""
    logo = profile.resolved_logo()
    logo_w = {"classic": 3.0, "modern": 2.8, "compact": 2.2}.get(style, 2.8)
    name_sz = {"classic": 13, "modern": 14, "compact": 12}.get(style, 13)
    tag_sz = {"classic": 10, "modern": 9, "compact": 9}.get(style, 9)

    if logo is not None:
        # Двухколоночная шапка: логотип слева, название/контакты справа
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left, right = table.rows[0].cells
        try:
            left.paragraphs[0].add_run().add_picture(str(logo), width=Cm(logo_w))
        except Exception:
            pass
        p = right.paragraphs[0]
        r = p.add_run(profile.name)
        _set_run_font(r, size=name_sz, bold=True, color=colors["name"])
        if profile.tagline:
            p2 = right.add_paragraph()
            r2 = p2.add_run(profile.tagline)
            _set_run_font(r2, size=tag_sz, color=colors["muted"])
        if style != "compact":
            _add_contact_lines(right, profile, colors, size=8 if style == "modern" else 9)
        else:
            # compact: контакты одной строкой под шапкой
            pass
        if style == "compact":
            _add_contact_paragraph(doc, profile, colors)
    else:
        p_lab = doc.add_paragraph()
        r = p_lab.add_run(profile.name)
        _set_run_font(r, size=name_sz, bold=True, color=colors["name"])
        if profile.tagline:
            p_sub = doc.add_paragraph(profile.tagline)
            for run in p_sub.runs:
                _set_run_font(run, size=tag_sz, color=colors["muted"])
        _add_contact_paragraph(doc, profile, colors)


def _add_contact_lines(cell, profile: LabProfile, colors: dict, *, size: int = 9) -> None:
    bits = [profile.address, profile.phone, profile.email, profile.accreditation, profile.website]
    line = " · ".join(b for b in bits if b and b.strip())
    if not line:
        return
    p = cell.add_paragraph()
    r = p.add_run(line)
    _set_run_font(r, size=size, color=colors["muted"])


def _add_contact_paragraph(doc: Document, profile: LabProfile, colors: dict) -> None:
    bits = [profile.address, profile.phone, profile.email, profile.accreditation, profile.website]
    line = " · ".join(b for b in bits if b and b.strip())
    if not line:
        return
    p = doc.add_paragraph(line)
    for run in p.runs:
        _set_run_font(run, size=9, color=colors["muted"])


def generate_kp_docx(
    proposal: CommercialProposal,
    output_path: Path | str,
    *,
    style: str | None = None,
    lab_profile: LabProfile | None = None,
) -> Path:
    """Формирует КП: марки и итоги, логотип, 3 стиля оформления."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    profile = lab_profile or load_lab_profile()
    style_key = (style or profile.kp_style or "classic").strip().lower()
    if style_key not in _STYLE_COLORS:
        style_key = "classic"
    colors = _STYLE_COLORS[style_key]

    doc = Document()
    section = doc.sections[0]
    if style_key == "compact":
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.5)
    else:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    _add_header_block(doc, profile, colors, style=style_key)
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    title_size = 14 if style_key == "compact" else 16
    _set_run_font(tr, size=title_size, bold=True, color=colors["title"])

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = date_p.add_run(proposal.created_at.strftime("%d.%m.%Y"))
    _set_run_font(dr, size=10, color=colors["muted"])

    doc.add_paragraph()

    intro = doc.add_paragraph(build_intro_text(proposal.subject, proposal.customer))
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in intro.runs:
        _set_run_font(run, size=11, bold=True)

    if proposal.note:
        note_p = doc.add_paragraph(proposal.note)
        note_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in note_p.runs:
            _set_run_font(run, size=10)

    doc.add_paragraph()
    lead = doc.add_paragraph(
        "Стоимость испытаний по маркам кабельной продукции (руб., без детализации по видам испытаний):"
    )
    for run in lead.runs:
        _set_run_font(run, size=10, color=colors["muted"])

    headers = (
        "№",
        "Марка кабельной продукции",
        "Без НДС",
        f"НДС {int(proposal.vat_rate * 100)}%",
        "С НДС",
    )
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], colors["header_bg"])
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for idx, line in enumerate(proposal.marks, start=1):
        row = table.add_row().cells
        values = (
            str(idx),
            line.mark,
            format_money(line.total_without_vat),
            format_money(line.vat_amount),
            format_money(line.total_with_vat),
        )
        for j, (cell, val) in enumerate(zip(row, values)):
            cell.text = val
            if j == 0:
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif j >= 2:
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            for paragraph in cell.paragraphs:
                paragraph.alignment = align
                for run in paragraph.runs:
                    _set_run_font(run, size=10)

    total_row = table.add_row().cells
    total_row[0].merge(total_row[1])
    total_row[0].text = "ИТОГО"
    total_row[2].text = format_money(proposal.total_without_vat)
    total_row[3].text = format_money(proposal.total_vat)
    total_row[4].text = format_money(proposal.total_with_vat)
    for i, cell in enumerate(total_row):
        _shade_cell(cell, colors["total_bg"])
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in paragraph.runs:
                _set_run_font(run, size=10, bold=True)

    doc.add_paragraph()

    validity = doc.add_paragraph(
        f"Предложение действительно в течение {proposal.validity_days} календарных дней "
        f"с даты составления."
    )
    for run in validity.runs:
        _set_run_font(run, size=10, color=colors["muted"])

    footer = doc.add_paragraph(
        "Срок выполнения работ и условия оплаты согласовываются дополнительно. "
        "Настоящее предложение не является публичной офертой."
    )
    for run in footer.runs:
        _set_run_font(run, size=9, color=RGBColor(0x94, 0xA3, 0xB8))

    if profile.name:
        sign = doc.add_paragraph()
        sign.add_run("\n")
        r = sign.add_run(profile.name)
        _set_run_font(r, size=10, bold=True, color=colors["name"])

    doc.save(str(path))
    return path.resolve()


def proposal_from_calculations(
    *,
    customer: str,
    subject: str,
    calculations: list[dict],
    note: str | None = None,
    validity_days: int = DEFAULT_VALIDITY_DAYS,
) -> CommercialProposal:
    """Собирает модель КП из записей calculations (summary dict)."""
    marks: list[KPMarkLine] = []
    vat_rate = 0.22
    for row in calculations:
        without = float(row["total_cost_without_vat"])
        with_vat = float(row["total_cost_with_vat"])
        rate = float(row.get("vat_rate") or 0.22)
        vat_rate = rate
        marks.append(
            KPMarkLine(
                mark=row["mark"],
                total_without_vat=without,
                vat_amount=round(with_vat - without, 2),
                total_with_vat=with_vat,
                calculation_id=row.get("id"),
            )
        )
    return CommercialProposal(
        customer=customer,
        subject=subject,
        note=note,
        marks=marks,
        vat_rate=vat_rate,
        validity_days=validity_days,
        created_at=datetime.now(),
    )


def generate_kp_from_db(
    *,
    customer: str,
    subject: str,
    calculation_ids: list[int],
    output_path: Path | str,
    db_path: Path | str | None = None,
    note: str | None = None,
    style: str | None = None,
) -> Path:
    """Загружает расчёты из БД и сохраняет КП в Word."""
    from ..persistence.sqlite_repo import (
        DB_PATH_DEFAULT,
        get_calculations_for_kp,
        update_calculation_output_path,
    )

    if db_path is None:
        db_path = DB_PATH_DEFAULT

    if not calculation_ids:
        raise ValueError("Выберите хотя бы один расчёт для КП")

    rows = get_calculations_for_kp(calculation_ids, db_path=db_path)
    if not rows:
        raise ValueError("Расчёты не найдены в БД")

    proposal = proposal_from_calculations(
        customer=customer,
        subject=subject,
        calculations=rows,
        note=note,
    )
    path = generate_kp_docx(proposal, output_path, style=style)
    proposal.output_path = str(path)
    for line in proposal.marks:
        if line.calculation_id:
            update_calculation_output_path(line.calculation_id, str(path), db_path)
    return path


def render_kp_style_previews(
    output_dir: Path | str | None = None,
) -> list[Path]:
    """Генерирует 3 образца бланка КП для выбора стиля (без БД)."""
    from ..config import GENERATED_DIR

    out = Path(output_dir) if output_dir else GENERATED_DIR / "kp_style_previews"
    out.mkdir(parents=True, exist_ok=True)
    proposal = CommercialProposal(
        customer="ООО «Пример Заказчик»",
        subject="Проведение приемосдаточных испытаний",
        note="Образец бланка — выберите стиль в data/lab_profile.yaml (kp_style).",
        marks=[
            KPMarkLine(
                mark="ВВГнг(А)-LS 3х2,5",
                total_without_vat=10000.0,
                vat_amount=2200.0,
                total_with_vat=12200.0,
            ),
            KPMarkLine(
                mark="КСБнг(А)-FRLS 4х1,5",
                total_without_vat=15000.0,
                vat_amount=3300.0,
                total_with_vat=18300.0,
            ),
        ],
        vat_rate=0.22,
        validity_days=30,
        created_at=datetime.now(),
    )
    profile = load_lab_profile()
    paths: list[Path] = []
    for style in ("classic", "modern", "compact"):
        path = out / f"КП_образец_{style}.docx"
        generate_kp_docx(proposal, path, style=style, lab_profile=profile)
        paths.append(path)
    return paths
