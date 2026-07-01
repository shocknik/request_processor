"""
Генерация коммерческого предложения (КП) в формате Word.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import CommercialProposal, KPMarkLine

LAB_NAME = "ООО «НИЦ Кабель-Тест»"
LAB_TAGLINE = "Испытательный центр кабельной продукции"
DEFAULT_VALIDITY_DAYS = 30
VAT_PERCENT = 22


def format_money(amount: float) -> str:
    """1 234 567,89"""
    whole, frac = f"{amount:.2f}".split(".")
    grouped = ""
    for i, ch in enumerate(reversed(whole)):
        if i and i % 3 == 0:
            grouped = " " + grouped
        grouped = ch + grouped
    return f"{grouped},{frac}"


def build_intro_text(subject: str, customer: str) -> str:
    subject = subject.strip() or "Проведение испытаний"
    customer = customer.strip()
    if customer:
        return f"{subject} для {customer}"
    return subject


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, *, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def generate_kp_docx(proposal: CommercialProposal, output_path: Path | str) -> Path:
    """Формирует КП без детализации строк расчёта — только марки и итоги."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Шапка организации
    p_lab = doc.add_paragraph()
    p_lab.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_lab.add_run(LAB_NAME)
    _set_run_font(r, size=13, bold=True, color=RGBColor(0x1E, 0x3A, 0x5F))
    p_sub = doc.add_paragraph(LAB_TAGLINE)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_sub.runs:
        _set_run_font(run, size=10, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    _set_run_font(tr, size=16, bold=True, color=RGBColor(0x25, 0x63, 0xEB))

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = date_p.add_run(proposal.created_at.strftime("%d.%m.%Y"))
    _set_run_font(dr, size=10, color=RGBColor(0x64, 0x74, 0x8B))

    doc.add_paragraph()

    intro = doc.add_paragraph(build_intro_text(proposal.subject, proposal.customer))
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in intro.runs:
        _set_run_font(run, size=11, bold=True)

    if proposal.note:
        note_p = doc.add_paragraph(proposal.note)
        note_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in note_p.runs:
            _set_run_font(run, size=10)

    doc.add_paragraph()
    lead = doc.add_paragraph(
        "Стоимость испытаний по маркам кабельной продукции (руб., без детализации по видам испытаний):"
    )
    for run in lead.runs:
        _set_run_font(run, size=10, color=RGBColor(0x47, 0x55, 0x69))

    # Таблица
    headers = (
        "№",
        "Марка кабельной продукции",
        "Без НДС",
        f"НДС {int(proposal.vat_rate * 100)}%",
        "С НДС",
    )
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], "2563EB")
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for idx, line in enumerate(proposal.marks, start=1):
        row = table.add_row().cells
        values = (
            str(idx),
            line.mark,
            format_money(line.total_without_vat),
            format_money(line.vat_amount),
            format_money(line.total_with_vat),
        )
        for j, (cell, val) in enumerate(zip(row, values)):
            cell.text = val
            if j == 0:
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif j >= 2:
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            for paragraph in cell.paragraphs:
                paragraph.alignment = align
                for run in paragraph.runs:
                    _set_run_font(run, size=10)

    # Итого
    total_row = table.add_row().cells
    total_row[0].merge(total_row[1])
    total_row[0].text = "ИТОГО"
    total_row[2].text = format_money(proposal.total_without_vat)
    total_row[3].text = format_money(proposal.total_vat)
    total_row[4].text = format_money(proposal.total_with_vat)
    for i, cell in enumerate(total_row):
        _shade_cell(cell, "E2E8F0")
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in paragraph.runs:
                _set_run_font(run, size=10, bold=True)

    doc.add_paragraph()

    validity = doc.add_paragraph(
        f"Предложение действительно в течение {proposal.validity_days} календарных дней "
        f"с даты составления."
    )
    for run in validity.runs:
        _set_run_font(run, size=10, color=RGBColor(0x64, 0x74, 0x8B))

    footer = doc.add_paragraph(
        "Срок выполнения работ и условия оплаты согласовываются дополнительно. "
        "Настоящее предложение не является публичной офертой."
    )
    for run in footer.runs:
        _set_run_font(run, size=9, color=RGBColor(0x94, 0xA3, 0xB8))

    doc.save(str(path))
    return path.resolve()


def proposal_from_calculations(
    *,
    customer: str,
    subject: str,
    calculations: list[dict],
    note: str | None = None,
    validity_days: int = DEFAULT_VALIDITY_DAYS,
) -> CommercialProposal:
    """Собирает модель КП из записей calculations (summary dict)."""
    marks: list[KPMarkLine] = []
    vat_rate = 0.22
    for row in calculations:
        without = float(row["total_cost_without_vat"])
        with_vat = float(row["total_cost_with_vat"])
        rate = float(row.get("vat_rate") or 0.22)
        vat_rate = rate
        marks.append(
            KPMarkLine(
                mark=row["mark"],
                total_without_vat=without,
                vat_amount=round(with_vat - without, 2),
                total_with_vat=with_vat,
                calculation_id=row.get("id"),
            )
        )
    return CommercialProposal(
        customer=customer,
        subject=subject,
        note=note,
        marks=marks,
        vat_rate=vat_rate,
        validity_days=validity_days,
        created_at=datetime.now(),
    )


def generate_kp_from_db(
    *,
    customer: str,
    subject: str,
    calculation_ids: list[int],
    output_path: Path | str,
    db_path: Path | str | None = None,
    note: str | None = None,
) -> Path:
    """Загружает расчёты из БД и сохраняет КП в Word."""
    from .sqlite_repo import DB_PATH_DEFAULT, get_calculations_for_kp, update_calculation_output_path

    if db_path is None:
        db_path = DB_PATH_DEFAULT

    if not calculation_ids:
        raise ValueError("Выберите хотя бы один расчёт для КП")

    rows = get_calculations_for_kp(calculation_ids, db_path=db_path)
    if not rows:
        raise ValueError("Расчёты не найдены в БД")

    proposal = proposal_from_calculations(
        customer=customer,
        subject=subject,
        calculations=rows,
        note=note,
    )
    path = generate_kp_docx(proposal, output_path)
    proposal.output_path = str(path)
    for line in proposal.marks:
        if line.calculation_id:
            update_calculation_output_path(line.calculation_id, str(path), db_path)
    return path