"""
Импорт строк таблицы приёмки ТУ → acceptance_items (волна 2, ТЗ v3).

Источник: локальный .docx (поколение A) или raw_text (fallback для «рамок»).
Файлы ТУ **не** коммитятся — только SQLite.

Пункты требований/методов **разворачиваются** в отдельные clause
(не храним «2.3.1-2.3.6» одной строкой — решение оператора).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ..config import PROJECT_ROOT
from ..persistence.sqlite_repo import (
    DB_PATH_DEFAULT,
    add_acceptance_item,
    get_connection,
    list_acceptance_items,
    upsert_norm_document,
)

# --- paths ---

TU_CORPUS = PROJECT_ROOT / "data" / "training" / "rag_corpus" / "tu"
RAW_TEXT_DIR = PROJECT_ROOT / "data" / "knowledge" / "manufacturer_v1" / "raw_text"

# Эталоны волны 2 (согласовано v3)
ETALON_GLOBS: tuple[tuple[str, str, str], ...] = (
    # (doc_id, filename_glob, preferred_source)
    ("ТУ 27.31.11-131-47273194-2025", "*131*2025*.docx", "docx"),
    ("ТУ 27.32.13-141-47273194-2024", "*141*2024*.docx", "docx"),
    ("ТУ 16.К99-005-01", "*005-01*.docx", "raw_text"),  # framed → raw_text methods
)

_RE_CLAUSE = re.compile(r"\d+(?:\.\d+){1,5}")
_RE_GROUP = re.compile(r"^[СCСсSs](\d+)$|^[ПPпp](\d+)$|^[ТTтt](\d+)$")
_RE_RANGE = re.compile(
    r"(\d+(?:\.\d+)*)\s*[\-–—÷]\s*(\d+(?:\.\d+)*)",
)
_RE_METHOD_LINE = re.compile(
    r"^\s*(\d+(?:\.\d+){1,4})\.?\s+"
    r"(.+?)"
    r"\s*\(([^)]{3,80})\)\s*"
    r"(?:проводят|определяют|выполняют)?\s*"
    r"(?:по\s+)?(ГОСТ[^\n.]{3,80})?",
    re.IGNORECASE,
)
_RE_TU_IN_NAME = re.compile(
    r"(ТУ\s*[\d.КкA-Za-zА-Яа-я\-–—]+)",
    re.IGNORECASE,
)

_NON_BILLABLE = re.compile(
    r"маркиров|упаков|срок\s+служб|комплектност|сопровод",
    re.IGNORECASE,
)


@dataclass
class ParsedAcceptanceRow:
    name_exact: str
    requirement_clauses: list[str] = field(default_factory=list)
    method_clauses: list[str] = field(default_factory=list)
    test_category: str | None = None
    group_code: str | None = None
    billable: bool = True
    method_external: list[dict[str, str]] = field(default_factory=list)
    notes: str = ""
    sort_order: int = 0


@dataclass
class ParsedAcceptanceCatalog:
    doc_id: str
    title: str
    source_path: str
    source_format: str  # docx_clean | raw_text | doc_framed
    items: list[ParsedAcceptanceRow] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()


def _cell_text(cell) -> str:
    return _norm_ws(" ".join(p.text for p in cell.paragraphs if p.text))


def _uniq_cells(cells: list[str]) -> list[str]:
    out: list[str] = []
    for c in cells:
        c = _norm_ws(c)
        if not out or out[-1] != c:
            out.append(c)
    return out


def expand_clause_refs(raw: str) -> list[str]:
    """
    «2.5.1, 2.5.2» / «2.3.1 - 2.3.6» / «2.3.7 2.3.13» → список отдельных пунктов.
    Диапазон разворачивается только если совпадает префикс (одинаковая глубина).
    """
    if not raw or not str(raw).strip():
        return []
    text = _norm_ws(str(raw)).replace("–", "-").replace("—", "-").replace("÷", "-")
    # убрать «таблица N, пункт K» хвосты — оставить clause-like
    text = re.sub(r"таблица\s*\d+[^\d]*", " ", text, flags=re.I)
    text = re.sub(r"пункт[аы]?\s*", " ", text, flags=re.I)
    text = re.sub(r"\([^)]*кроме[^)]*\)", " ", text, flags=re.I)

    parts: list[str] = []
    # split by comma or semicolon first
    chunks = re.split(r"[,;]+", text)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        m = _RE_RANGE.search(chunk)
        if m and chunk.count("-") <= 2:
            a, b = m.group(1), m.group(2)
            expanded = _expand_numeric_range(a, b)
            if expanded:
                parts.extend(expanded)
                continue
        # space-separated clauses
        found = _RE_CLAUSE.findall(chunk)
        if found:
            parts.extend(found)
        elif _RE_CLAUSE.fullmatch(chunk.replace(" ", "")):
            parts.append(chunk.replace(" ", ""))

    # unique preserve order
    seen: set[str] = set()
    out: list[str] = []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out


def _expand_numeric_range(a: str, b: str) -> list[str] | None:
    pa, pb = a.split("."), b.split(".")
    if len(pa) != len(pb):
        return None
    if pa[:-1] != pb[:-1]:
        return None
    try:
        start, end = int(pa[-1]), int(pb[-1])
    except ValueError:
        return None
    if end < start or end - start > 40:
        return None
    prefix = ".".join(pa[:-1])
    if prefix:
        return [f"{prefix}.{i}" for i in range(start, end + 1)]
    return [str(i) for i in range(start, end + 1)]


def _category_from_group(group: str | None) -> str | None:
    if not group:
        return None
    g = group.strip().upper().replace("C", "С").replace("P", "П").replace("T", "Т")
    if g.startswith("С") or g.startswith("C"):
        return "psi"
    if g.startswith("П") or g.startswith("P"):
        return "periodic"
    if g.startswith("Т") or g.startswith("T"):
        return "type"
    return None


def _guess_billable(name: str) -> bool:
    return not bool(_NON_BILLABLE.search(name or ""))


def _is_acceptance_header(cells: list[str]) -> bool:
    j = " ".join(cells).lower()
    has_view = "вид испытан" in j or "вид провер" in j or "наименование" in j
    has_pts = "требован" in j or "пункт" in j or "метод" in j
    return has_view and has_pts


def _is_data_row(cells: list[str]) -> bool:
    if len(cells) < 3:
        return False
    j = " ".join(cells).lower()
    if "группа испытан" in j and "вид" in j:
        return False
    if "технических требован" in j and "методов" in j and len(cells) <= 4:
        # sub-header
        return False
    # need a name-like cell and at least one clause
    blob = " ".join(cells)
    if not _RE_CLAUSE.search(blob):
        return False
    return True


def _map_row(cells: list[str], *, category_hint: str | None, sort_order: int) -> ParsedAcceptanceRow | None:
    cells = [c for c in cells if c is not None]
    if len(cells) < 3:
        return None
    group: str | None = None
    name = ""
    req_raw = ""
    meth_raw = ""

    # Typical: [С1, name, req, method] or [С1, name, req] with method in 4th after uniq
    c0 = cells[0]
    if re.match(r"^[СCСсSsПPпpТTтt]\d+", c0.replace(" ", "")):
        group = re.sub(r"\s+", "", c0.upper().replace("C", "С").replace("P", "П").replace("T", "Т"))
        # normalize latin
        group = group[0] + re.sub(r"\D", "", group[1:] or "0")
        if len(cells) >= 4:
            name, req_raw, meth_raw = cells[1], cells[2], cells[3]
        elif len(cells) == 3:
            name, req_raw = cells[1], cells[2]
            # sometimes method glued
            parts = re.split(r"\s{2,}|\t", req_raw)
            if len(parts) >= 2 and _RE_CLAUSE.search(parts[-1]):
                meth_raw = parts[-1]
                req_raw = parts[0]
        else:
            return None
    else:
        # no group column
        if len(cells) >= 3:
            name, req_raw, meth_raw = cells[0], cells[1], cells[2]
        else:
            return None

    name = _norm_ws(name)
    if len(name) < 4:
        return None
    # skip pure headers / приложение А (ссылочные ГОСТ)
    nl = name.lower()
    if "вид испытания" in nl and len(name) < 40:
        return None
    if re.match(r"^(гост|gost|iec)\b", nl):
        return None
    if nl.startswith("испытания проводят") or nl.startswith("правила прием"):
        return None

    reqs = expand_clause_refs(req_raw)
    meths = expand_clause_refs(meth_raw)
    if not reqs and not meths:
        return None

    cat = _category_from_group(group) or category_hint
    return ParsedAcceptanceRow(
        name_exact=name[:300],
        requirement_clauses=reqs,
        method_clauses=meths,
        test_category=cat,
        group_code=group,  # optional field; stored if present
        billable=_guess_billable(name),
        sort_order=sort_order,
        notes=f"import: req_raw={req_raw[:80]!r}" if ("-" in (req_raw or "") or "–" in (req_raw or "")) else "",
    )


def _category_hint_near_table(paragraphs: list[str], table_index: int) -> str | None:
    # look at last few captions
    blob = " ".join(paragraphs[-15:]).lower()
    if "периодическ" in blob:
        return "periodic"
    if "приёмо" in blob or "приемо" in blob or "сдаточн" in blob:
        return "psi"
    if "типов" in blob:
        return "type"
    return None


def parse_acceptance_from_docx(path: Path | str) -> ParsedAcceptanceCatalog:
    """Парсит таблицы «Состав … испытаний» из современного docx."""
    from docx import Document

    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(str(file_path))
    if file_path.suffix.lower() != ".docx":
        raise ValueError(f"Ожидается .docx, получено: {file_path.suffix}")

    doc = Document(str(file_path))
    doc_id = _doc_id_from_filename(file_path.name)
    title = _title_from_doc(doc) or file_path.stem[:120]

    # map body elements order: collect para texts before each table roughly
    # simpler: for each table, use document paragraphs scan is weak;
    # use prev tables + full para list with table captions
    all_paras = [_norm_ws(p.text) for p in doc.paragraphs if p.text and p.text.strip()]

    items: list[ParsedAcceptanceRow] = []
    warnings: list[str] = []
    sort = 0
    tables_used = 0

    for ti, table in enumerate(doc.tables):
        rows_cells: list[list[str]] = []
        for row in table.rows:
            raw = [_cell_text(c) for c in row.cells]
            rows_cells.append(_uniq_cells(raw))

        if not rows_cells:
            continue
        header_hits = sum(1 for r in rows_cells[:3] if _is_acceptance_header(r))
        # also detect by first data-like after header keywords in any row
        joined_head = " ".join(" ".join(r) for r in rows_cells[:3]).lower()
        looks = (
            header_hits > 0
            or (
                "вид испытан" in joined_head
                and ("требован" in joined_head or "метод" in joined_head)
            )
            or (
                "группа испытан" in joined_head
                and "вид" in joined_head
                and "пункт" in joined_head
            )
        )
        if not looks:
            continue

        # caption hint: search paras containing "таблица" near index
        cat_hint = None
        for p in all_paras:
            pl = p.lower()
            if f"таблица {ti}" in pl or "состав при" in pl or "состав период" in pl:
                if "периодическ" in pl:
                    cat_hint = "periodic"
                elif "приём" in pl or "прием" in pl or "сдаточн" in pl:
                    cat_hint = "psi"
                elif "типов" in pl:
                    cat_hint = "type"
        # stronger: caption lines
        for p in all_paras:
            pl = p.lower()
            if "состав при" in pl and ("сдаточ" in pl or "приём" in pl or "прием" in pl):
                # will apply to following psi tables — set default if unset
                pass

        # Infer category from group codes in body if caption weak
        local_items: list[ParsedAcceptanceRow] = []
        for r in rows_cells:
            if not _is_data_row(r):
                continue
            # category from group in row preferred
            sort += 10
            parsed = _map_row(r, category_hint=cat_hint, sort_order=sort)
            if parsed:
                local_items.append(parsed)

        if not local_items:
            continue

        # if still no category, derive majority from groups
        if not cat_hint:
            cats = [i.test_category for i in local_items if i.test_category]
            if cats:
                cat_hint = max(set(cats), key=cats.count)
        for it in local_items:
            if not it.test_category:
                it.test_category = cat_hint

        items.extend(local_items)
        tables_used += 1

    if tables_used == 0:
        warnings.append("Не найдено таблиц приёмки (заголовки «Вид испытания» + пункты)")
    if not items:
        warnings.append("0 строк после разбора")

    # dedupe by name_norm + category
    items = _dedupe_items(items)

    return ParsedAcceptanceCatalog(
        doc_id=doc_id,
        title=title,
        source_path=str(file_path),
        source_format="docx_clean",
        items=items,
        warnings=warnings,
    )


def parse_acceptance_from_raw_text_methods(
    path: Path | str,
    *,
    doc_id: str | None = None,
) -> ParsedAcceptanceCatalog:
    """
    Fallback для framed ТУ (005): строки вида
    «4.3.1 Определение … (1.4.1, …) проводят по ГОСТ …»
    """
    file_path = Path(path)
    text = file_path.read_text(encoding="utf-8", errors="replace")
    # de-dup repeated frame blocks: keep unique lines order
    seen_lines: set[str] = set()
    unique_lines: list[str] = []
    for line in text.splitlines():
        ln = _norm_ws(line)
        if len(ln) < 12:
            continue
        key = ln[:120].lower()
        if key in seen_lines:
            continue
        seen_lines.add(key)
        unique_lines.append(ln)

    items: list[ParsedAcceptanceRow] = []
    sort = 0
    for ln in unique_lines:
        m = _RE_METHOD_LINE.match(ln)
        if not m:
            continue
        method_clause, name_part, req_blob, gost = m.group(1), m.group(2), m.group(3), m.group(4)
        name = _norm_ws(name_part)
        # skip non-test prose
        if not re.search(
            r"испытан|провер|измерен|определен|сопротив|затухан|напряжен|"
            r"герметич|изгиб|растяж|маркир|длин|волнов|экран|отражен",
            name,
            re.I,
        ):
            continue
        if re.match(r"^(гост|gost)\b", name.lower()):
            continue
        if name.lower().startswith("испытания проводят"):
            continue
        # слишком короткое «Маркировка» без контекста — шум raw_text
        if len(name) < 18 and not re.search(
            r"сопротив|затухан|напряжен|герметич|изгиб|растяж|волнов",
            name,
            re.I,
        ):
            continue
        reqs = expand_clause_refs(req_blob)
        if not reqs:
            reqs = expand_clause_refs(ln)
        sort += 10
        ext = []
        if gost:
            g = _norm_ws(gost)
            # обрезать по «и внешним» / лишним хвостам
            g = re.split(r"\s+и\s+внешн", g, maxsplit=1, flags=re.I)[0]
            g = re.sub(r"\s*\(\d+$", "", g)  # обрезанная скобка
            g = g.rstrip(".,; ")
            # «ГОСТ Р 58416 (8.3.1)» целиком если скобка закрыта
            gm = re.match(
                r"(ГОСТ[^.]{3,60}?)(?:\s*\(([^)]+)\))?\s*$",
                g,
                re.I,
            )
            if gm:
                ext.append(
                    {
                        "ext_doc_id": gm.group(1).strip()[:120],
                        "ext_clause_or_method": (gm.group(2) or "").strip(),
                    }
                )
            elif g:
                ext.append({"ext_doc_id": g[:120], "ext_clause_or_method": ""})
        items.append(
            ParsedAcceptanceRow(
                name_exact=name[:300],
                requirement_clauses=reqs[:12],
                method_clauses=[method_clause],
                test_category=None,  # unknown in methods dump
                group_code=None,
                billable=_guess_billable(name),
                method_external=ext,
                sort_order=sort,
                notes="import: raw_text methods fallback",
            )
        )

    items = _dedupe_items(items)
    did = doc_id or _doc_id_from_filename(file_path.name)
    return ParsedAcceptanceCatalog(
        doc_id=did,
        title=did,
        source_path=str(file_path),
        source_format="raw_text",
        items=items,
        warnings=[]
        if items
        else ["raw_text methods: 0 строк (framed table not recoverable)"],
    )


def _dedupe_items(items: list[ParsedAcceptanceRow]) -> list[ParsedAcceptanceRow]:
    seen: set[str] = set()
    out: list[ParsedAcceptanceRow] = []
    for it in items:
        key = f"{_norm_ws(it.name_exact).lower()}|{it.test_category or ''}"
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


def _doc_id_from_filename(name: str) -> str:
    stem = Path(name).stem
    stem = stem.replace("–", "-").replace("—", "-")
    # drop trailing «от date» / «с изм…»
    stem = re.sub(r"\s+от\s+\d.*$", "", stem, flags=re.I)
    stem = re.sub(r"\s+с\s+изм.*$", "", stem, flags=re.I)
    stem = re.sub(r"\s+изм\..*$", "", stem, flags=re.I)
    stem = _norm_ws(stem)
    # fix common typo 472731194 → keep as in file if present; normalize spaces
    if not stem.upper().startswith("ТУ"):
        stem = f"ТУ {stem}"
    # canonical space after ТУ
    stem = re.sub(r"^ТУ\s*", "ТУ ", stem, flags=re.I)
    return stem[:120]


def _title_from_doc(doc) -> str:
    for p in doc.paragraphs[:40]:
        t = _norm_ws(p.text)
        if 15 < len(t) < 180 and not t.lower().startswith("таблица"):
            if re.search(r"кабел|провод|шнур|ТУ", t, re.I):
                return t[:200]
    return ""


def resolve_etalon_path(glob_pat: str, *, corpus: Path = TU_CORPUS) -> Path | None:
    hits = sorted(corpus.glob(glob_pat))
    return hits[0] if hits else None


def clear_acceptance_for_document(
    norm_document_id: int,
    *,
    db_path: str | Path = DB_PATH_DEFAULT,
) -> int:
    """Удаляет старые acceptance_items документа (перед полным re-import)."""
    with get_connection(db_path) as conn:
        cur = conn.execute(
            "DELETE FROM acceptance_items WHERE norm_document_id = ?",
            (norm_document_id,),
        )
        return int(cur.rowcount or 0)


def import_parsed_catalog(
    catalog: ParsedAcceptanceCatalog,
    *,
    db_path: str | Path = DB_PATH_DEFAULT,
    replace: bool = True,
    manufacturer_hint: str | None = "ООО НПП Спецкабель",
) -> dict[str, Any]:
    """Пишет ParsedAcceptanceCatalog в SQLite."""
    nd_id = upsert_norm_document(
        catalog.doc_id,
        catalog.title,
        kind="tu",
        file_path=catalog.source_path if "rag_corpus" not in catalog.source_path.replace("\\", "/") else catalog.source_path,
        notes=f"acceptance import wave2; warnings={len(catalog.warnings)}",
        source_format=catalog.source_format,
        manufacturer_hint=manufacturer_hint,
        status="draft",
        db_path=db_path,
    )
    # Prefer not storing full path that might leak? It's local-only DB — ok for operator.
    deleted = 0
    if replace:
        deleted = clear_acceptance_for_document(nd_id, db_path=db_path)

    added = 0
    for it in catalog.items:
        add_acceptance_item(
            norm_document_id=nd_id,
            name_exact=it.name_exact,
            requirement_clauses=it.requirement_clauses,
            method_clauses=it.method_clauses,
            test_category=it.test_category,
            group_code=it.group_code,
            billable=it.billable,
            sort_order=it.sort_order,
            notes=it.notes or None,
            method_external=it.method_external or None,
            status="draft",
            db_path=db_path,
        )
        added += 1

    return {
        "doc_id": catalog.doc_id,
        "norm_document_id": nd_id,
        "items": added,
        "deleted_before": deleted,
        "warnings": list(catalog.warnings),
        "source_format": catalog.source_format,
        "source_path": catalog.source_path,
    }


def import_acceptance_docx(
    path: Path | str,
    *,
    db_path: str | Path = DB_PATH_DEFAULT,
    replace: bool = True,
    doc_id: str | None = None,
) -> dict[str, Any]:
    catalog = parse_acceptance_from_docx(path)
    if doc_id:
        catalog.doc_id = doc_id
    return import_parsed_catalog(catalog, db_path=db_path, replace=replace)


def import_acceptance_raw_text(
    path: Path | str,
    *,
    db_path: str | Path = DB_PATH_DEFAULT,
    replace: bool = True,
    doc_id: str | None = None,
) -> dict[str, Any]:
    catalog = parse_acceptance_from_raw_text_methods(path, doc_id=doc_id)
    return import_parsed_catalog(catalog, db_path=db_path, replace=replace)


def import_etalon_batch(
    *,
    db_path: str | Path = DB_PATH_DEFAULT,
    corpus: Path = TU_CORPUS,
    raw_dir: Path = RAW_TEXT_DIR,
    replace: bool = True,
) -> list[dict[str, Any]]:
    """Импорт согласованных эталонов 131, 141, 005."""
    results: list[dict[str, Any]] = []
    for doc_id, glob_pat, preferred in ETALON_GLOBS:
        if preferred == "docx":
            path = resolve_etalon_path(glob_pat, corpus=corpus)
            if not path:
                results.append(
                    {
                        "doc_id": doc_id,
                        "error": f"файл не найден: {glob_pat} в {corpus}",
                        "items": 0,
                    }
                )
                continue
            try:
                cat = parse_acceptance_from_docx(path)
                cat.doc_id = doc_id  # canonical id from TZ
                r = import_parsed_catalog(cat, db_path=db_path, replace=replace)
                results.append(r)
            except Exception as e:
                results.append({"doc_id": doc_id, "error": str(e), "items": 0})
        else:
            # raw_text preferred for 005
            raw_hits = sorted(raw_dir.glob("*005-01*.txt")) + sorted(
                raw_dir.glob("*005*.txt")
            )
            # also try docx parse first; if 0 items fall back
            docx_path = resolve_etalon_path(glob_pat, corpus=corpus)
            used = None
            if docx_path:
                try:
                    cat = parse_acceptance_from_docx(docx_path)
                    if cat.items:
                        cat.doc_id = doc_id
                        used = import_parsed_catalog(
                            cat, db_path=db_path, replace=replace
                        )
                        used["note"] = "docx produced items"
                except Exception:
                    used = None
            if used is None or used.get("items", 0) == 0:
                if not raw_hits:
                    results.append(
                        {
                            "doc_id": doc_id,
                            "error": "нет raw_text 005 и docx пуст",
                            "items": 0,
                        }
                    )
                    continue
                cat = parse_acceptance_from_raw_text_methods(
                    raw_hits[0], doc_id=doc_id
                )
                used = import_parsed_catalog(cat, db_path=db_path, replace=replace)
                used["note"] = f"raw_text fallback: {raw_hits[0].name}"
            results.append(used)
    return results


def try_match_prices(
    *,
    doc_id: str | None = None,
    db_path: str | Path = DB_PATH_DEFAULT,
) -> dict[str, Any]:
    """Опционально: проставить price_test_code через program_price_matcher."""
    from ..mapping.program_price_matcher import resolve_program_item_price_code

    items = list_acceptance_items(doc_id=doc_id, db_path=db_path, limit=5000)
    matched = 0
    with get_connection(db_path) as conn:
        for it in items:
            if it.get("price_test_code"):
                continue
            if not it.get("billable"):
                continue
            name = it.get("name_exact") or ""
            hit = resolve_program_item_price_code(name, db_path=db_path)
            code = hit.code if hit else None
            if code:
                conn.execute(
                    "UPDATE acceptance_items SET price_test_code = ? WHERE id = ?",
                    (str(code), int(it["id"])),
                )
                matched += 1
    return {"scanned": len(items), "matched": matched}
