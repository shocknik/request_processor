"""
Импорт программы испытаний из Word (.docx) — один документ за раз.

Учитывает реальные ПМИ Спецкабель и аналоги:
  - шапка «П Р О Г Р А М М А» + абзацы объекта/цели;
  - **одна** марка в тексте («марки X …») или **список** «1) МАРКА – описание»;
  - таблицы испытаний (3–N колонок, двойной заголовок, «продолжение таблицы»);
  - иногда испытания только в тексте (п. 2.4.1 …).

Не пакетный импорт. Только .docx.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class ParsedProgramItem:
    sort_order: int
    name: str
    requirement_doc: str = ""
    requirement_clause: str = ""
    method_doc: str = ""
    method_clause: str = ""
    applies_to: str = ""  # «1, 2, 3» — применяемость к маркам списка


@dataclass
class ParsedProgram:
    name: str
    test_type: str = ""
    cable_mark_text: str = ""  # одна или несколько марок (по строке)
    cable_marks: list[str] = field(default_factory=list)
    tu_ref: str = ""
    source_path: str = ""
    notes: str = ""
    items: list[ParsedProgramItem] = field(default_factory=list)


# --- regex ---

# Только настоящее «ТУ 16.…» / «ТУ 27.…», не «турированных»
_RE_TU = re.compile(
    r"(?<![А-Яа-яA-Za-z])(ТУ\s+\d[\w.КкKkА-Яа-я\-–—]*)",
    re.IGNORECASE,
)
_RE_TYPE = re.compile(
    r"(при[её]мо\-?сдаточн\w*|при[её]мочн\w*|периодическ\w*|исследовател\w*|"
    r"типов\w*|сертификацион\w*|контрольн\w*)\s+испытан",
    re.IGNORECASE,
)
# «марки XXX» до запятой / «изготовлен» / конца
_RE_MARKI = re.compile(
    r"марки\s+(.+?)(?=,\s*изготовлен|\s+изготовлен|\s+на\s+соответствие|"
    r"\s+по\s+ТУ|\s+образц|\.\s+[А-ЯA-Z]|\n|$)",
    re.IGNORECASE | re.DOTALL,
)
# нумерованный список марок: 1) МАРКА – описание
_RE_LIST_MARK = re.compile(
    r"^\s*(?:\d+[\)\.]|\–|\-|•)\s*"
    r"(.+?)"  # mark + optional description
    r"\s*$",
)
# отделить обозначение от «– кабель …»
_RE_MARK_DESC_SPLIT = re.compile(r"\s+[–—\-]\s+(?=кабел|провод|шнур)", re.IGNORECASE)
_RE_NUM_START = re.compile(r"^(\d+)[\.\)]\s*(.+)$", re.DOTALL)
# пункт метода/испытания в тексте: 2.4.1 Испытания на …
_RE_TEXT_ITEM = re.compile(
    r"^\s*(\d+(?:\.\d+){1,3})\.?\s+"
    r"((?:Испытан|Проверк|Определен|Контрол|Измерен|Стойкость|Нераспространение|"
    r"Проверка|Огнестойк)[^\n]{5,180})",
    re.IGNORECASE,
)
# эвристика «похоже на марку»
_RE_MARKISH = re.compile(
    r"(нг\s*\([АA]\)|Cat\s*\d|F/?UTP|S/?FTP|х\d|x\d|×\d|\d+[xх]\d|"
    r"FRLS|FRHF|HF|LS|ХЛ|СКАБ|СПЕЦЛАН|ВВГ|КГВВ|ПвБ)",
    re.IGNORECASE,
)

_HEADER_KEYS = (
    "вид испытаний",
    "наименование проверяемого",
    "наименование параметра",
    "пункт пми",
    "пункты ту",
    "технических требований",
    "методов испытаний",
    "применяемость",
    "№ п/п",
    "n п/п",
    "объем выборки",
    "место испытаний",
)


def _cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def _uniq_cells(cells: list[str]) -> list[str]:
    """Убрать дубликаты от merge ячеек Word (одинаковые подряд)."""
    out: list[str] = []
    for c in cells:
        c = re.sub(r"\s+", " ", (c or "").replace("\xa0", " ")).strip()
        if not out or out[-1] != c:
            out.append(c)
    return out


def _normalize_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()


def _is_header_row(cells: list[str]) -> bool:
    joined = " ".join(cells).lower()
    hits = sum(1 for k in _HEADER_KEYS if k in joined)
    return hits >= 1 and not re.match(r"^\d", cells[0] if cells else "")


def _is_section_row(cells: list[str]) -> bool:
    """Строка-раздел: «КОНСТРУКЦИЯ», «ЭЛЕКТРИЧЕСКИЕ ПАРАМЕТРЫ»."""
    nonempty = [c for c in cells if c.strip()]
    if len(nonempty) != 1:
        return False
    t = nonempty[0].strip()
    if len(t) < 3 or len(t) > 60:
        return False
    if re.match(r"^\d", t):
        return False
    letters = [ch for ch in t if ch.isalpha()]
    if not letters:
        return False
    upper = sum(1 for ch in letters if ch.upper() == ch)
    return upper / len(letters) > 0.7


def _detect_columns(header_cells: list[str]) -> dict[str, int]:
    """Индексы колонок по заголовку."""
    roles: dict[str, int] = {}
    for i, h in enumerate(header_cells):
        hl = h.lower()
        if "№" in h or hl.startswith("n ") or "п/п" in hl or "пункт пми" in hl:
            roles.setdefault("num", i)
        elif "вид" in hl or "наименование" in hl or "параметр" in hl:
            roles.setdefault("name", i)
        elif "техническ" in hl and "треб" in hl:
            roles.setdefault("req", i)
        elif ("пункт" in hl or "пункты" in hl) and "треб" in hl:
            roles.setdefault("req", i)
        elif "метод" in hl:
            roles.setdefault("meth", i)
        elif ("пункт" in hl or "пункты" in hl) and "ту" in hl and "треб" not in hl:
            # «Пункты ТУ …» spanning — often requirements overall
            roles.setdefault("req_or_tu", i)
        elif "применяем" in hl or "объект" in hl:
            roles.setdefault("applies", i)
        elif "требован" in hl and "нд" in hl:
            roles.setdefault("req", i)
        elif "правил" in hl or "iec" in hl:
            roles.setdefault("meth_nd", i)
        elif re.fullmatch(r"пункт\s*ту", hl) or (hl.startswith("пункт ту")):
            roles.setdefault("req", i)
    # fallback classic 4-col
    if "name" not in roles and len(header_cells) >= 2:
        roles["name"] = 1
    if "num" not in roles and len(header_cells) >= 1:
        roles["num"] = 0
    if "req" not in roles:
        if "req_or_tu" in roles:
            roles["req"] = roles["req_or_tu"]
        elif len(header_cells) >= 3:
            roles["req"] = 2
    if "meth" not in roles and len(header_cells) >= 4:
        roles["meth"] = 3
    return roles


def _split_doc_and_clause(text: str, default_doc: str = "") -> tuple[str, str]:
    text = _normalize_ws(text)
    if not text:
        return default_doc, ""
    m = _RE_TU.search(text)
    if m:
        doc = m.group(1).replace("–", "-").replace("—", "-")
        clause = (text[: m.start()] + text[m.end() :]).strip(" ,;()")
        # «ТУ … (1.4.1)» 
        paren = re.search(r"\(([^)]+)\)\s*$", text)
        if paren and re.search(r"\d", paren.group(1)):
            clause = paren.group(1).strip()
        return doc, clause or text
    if re.match(r"^[\d\s,.\-–—;÷+()табл.п]+$", text, re.I):
        return default_doc, text.replace("–", "-").replace("÷", "-")
    return default_doc, text


def _clause_tokens(clause: str) -> list[str]:
    """Разбить «1.3.3 1.3.11» или «4.2.2 4.2.3» на токены."""
    if not clause:
        return []
    parts = re.split(r"[\s,;]+", clause.replace("÷", "-"))
    return [p for p in parts if re.search(r"\d", p)]


def _split_jammed_test_name(name: str) -> list[str]:
    """
    «Проверка A Проверка B» / «Проверка A Проверку B» / «Определение X Определение Y».
    """
    name = _normalize_ws(name)
    # stem groups (case-insensitive)
    stems = (
        r"Проверк\w*",
        r"Определен\w*",
        r"Испытан\w*",
        r"Контрол\w*",
        r"Измерен\w*",
    )
    for stem in stems:
        pat = re.compile(
            rf"^({stem}\s.+?)(?=\s+(?:{stem})\s)",
            re.IGNORECASE | re.DOTALL,
        )
        m = pat.match(name)
        if m and len(m.group(1).strip()) >= 8:
            first = m.group(1).strip()
            rest = name[m.end() :].strip()
            if len(rest) >= 8:
                return [first, rest]
    return [name]


def _is_strict_mark(m: str) -> bool:
    """Отсев мусора от find_cable_marks / обрывков фраз."""
    m = _normalize_ws(m)
    if len(m) < 6 or len(m) > 100:
        return False
    if not re.search(r"[A-Za-zА-Яа-яЁё]", m):
        return False
    if not re.search(r"\d", m):
        return False
    low = m.lower()
    bad_bits = (
        "длин",
        "образ",
        "номер",
        "пункт",
        "табл",
        "менее",
        "более",
        "также номер",
        "приведен",
        "соответств",
        "требован",
    )
    if any(b in low for b in bad_bits):
        return False
    # «0,52» / «25 м» 
    if re.fullmatch(r"[\d,\.\s]+м?", m, re.I):
        return False
    return bool(_RE_MARKISH.search(m))


def _clean_mark_token(raw: str) -> str:
    """Оставить обозначение марки без длинного описания."""
    raw = _normalize_ws(raw)
    if not raw:
        return ""
    # strip leading list junk
    raw = re.sub(r"^\d+[\)\.]\s*", "", raw)
    parts = _RE_MARK_DESC_SPLIT.split(raw, maxsplit=1)
    mark = parts[0].strip(" ,;.")
    # cut trailing «в соответствии…»
    mark = re.split(r"\s+в\s+соответствии\s+", mark, maxsplit=1, flags=re.I)[0]
    mark = re.split(r"\s+на\s+соответствие\s+", mark, maxsplit=1, flags=re.I)[0]
    # limit length
    if len(mark) > 120:
        # try find_cable_marks
        mark = mark[:120]
    return mark.strip()


def _extract_marks_from_paragraphs(paragraphs: list[str]) -> list[str]:
    marks: list[str] = []
    seen: set[str] = set()
    from_list = 0

    def add(m: str, *, from_list_item: bool = False) -> None:
        nonlocal from_list
        m = _clean_mark_token(m)
        if not _is_strict_mark(m):
            return
        key = m.lower()
        if key in seen:
            return
        if m.lower().startswith(("объект", "цель", "программа", "таблица")):
            return
        if m.endswith(".") and len(m) > 80:
            return
        seen.add(key)
        marks.append(m)
        if from_list_item:
            from_list += 1

    text = "\n".join(paragraphs)

    # 1) numbered list after «марки кабеля» / «типового представителя» — приоритет
    list_mode = False
    for p in paragraphs:
        pl = p.lower()
        if re.search(
            r"выбран[ыа]?\s+марк|марки\s+кабел|типового\s+представител", pl
        ):
            list_mode = True
            if ":" in p:
                after = p.split(":", 1)[1].strip()
                if after and not re.match(r"^\d+[\)\.]", after):
                    if _RE_MARKISH.search(after):
                        add(after, from_list_item=True)
            continue
        if list_mode:
            mlist = _RE_LIST_MARK.match(p)
            if mlist:
                add(mlist.group(1), from_list_item=True)
                continue
            if re.match(r"^\d+\.\s+(Объект|Цель|Программа|Образцы)", p, re.I):
                list_mode = False
            elif re.match(r"^(Образцы|2\.\s|3\.\s)", p, re.I):
                list_mode = False
            elif len(p) < 40 and not _RE_MARKISH.search(p) and re.match(
                r"^\d+\.\s+", p
            ):
                list_mode = False

    # 2) «марки XXX» в одной фразе (если списка мало)
    if from_list < 2:
        for m in _RE_MARKI.finditer(text):
            chunk = _normalize_ws(m.group(1))
            if re.search(r"\d+\)", chunk):
                continue
            pieces = re.split(r"\s*,\s*(?=[A-Za-zА-Яа-я])|\s+и\s+(?=[A-Za-zА-Яа-я])", chunk)
            if len(pieces) > 1 and all(
                _RE_MARKISH.search(p) and len(p) < 80 for p in pieces[:4]
            ):
                for p in pieces:
                    add(p)
            else:
                add(chunk)

    # 3) find_cable_marks — только если почти ничего не нашли
    if len(marks) < 1:
        try:
            from ..extraction.pdf_extractor import find_cable_marks

            blob = "\n".join(paragraphs[:40])
            for match in find_cable_marks(blob):
                add(match.mark)
        except Exception:
            pass

    # 4) «выбрана марка X – описание» одной строкой
    if not marks:
        for p in paragraphs:
            m = re.search(
                r"выбран[аоы]?\s+марка\s+(.+?)(?:\s+[–—\-]\s+|$)",
                p,
                re.I,
            )
            if m:
                add(m.group(1))

    return marks


def _parse_header(paragraphs: list[str]) -> dict[str, Any]:
    text = "\n".join(paragraphs)
    name = "Программа испытаний"
    for p in paragraphs:
        pl = p.lower()
        if "программ" in pl and "испытан" in pl and len(p) > 15:
            name = _normalize_ws(p)[:240]
            break
        if re.search(r"испытан\w+\s+кабел", pl) and len(p) > 25:
            name = _normalize_ws(p)[:240]
            break
        # spaced «П Р О Г Р А М М А» next lines
        if re.fullmatch(r"(?:П\s*){3,}А", p.replace(" ", ""), re.I) or re.search(
            r"п\s*р\s*о\s*г\s*р\s*а\s*м\s*м\s*а", p, re.I
        ):
            continue

    test_type = ""
    mtype = _RE_TYPE.search(text)
    if mtype:
        test_type = mtype.group(1)

    marks = _extract_marks_from_paragraphs(paragraphs)
    cable_mark_text = "\n".join(marks)

    tu_refs = []
    for m in _RE_TU.finditer(text):
        tu_refs.append(m.group(1).replace("–", "-").replace("—", "-"))
    # unique preserve order
    tu_ref = ", ".join(dict.fromkeys(tu_refs))

    return {
        "name": name,
        "test_type": test_type,
        "cable_marks": marks,
        "cable_mark_text": cable_mark_text,
        "tu_ref": tu_ref,
    }


def _parse_table_items(doc, default_tu: str) -> list[ParsedProgramItem]:
    items: list[ParsedProgramItem] = []
    auto_order = 0

    for table in doc.tables:
        if not table.rows or len(table.rows) < 2:
            continue
        # build raw rows
        raw_rows: list[list[str]] = []
        for row in table.rows:
            raw_rows.append(_uniq_cells([_cell_text(c) for c in row.cells]))

        # find header: up to first 3 rows
        roles: dict[str, int] | None = None
        data_start = 0
        for hi in range(min(3, len(raw_rows))):
            if _is_header_row(raw_rows[hi]):
                roles = _detect_columns(raw_rows[hi])
                data_start = hi + 1
                # if next also header (double header), merge roles and skip
                if hi + 1 < len(raw_rows) and _is_header_row(raw_rows[hi + 1]):
                    roles2 = _detect_columns(raw_rows[hi + 1])
                    roles = {**roles, **roles2}
                    data_start = hi + 2
                break
        if roles is None:
            # try classic if first data-like row has number
            if re.match(r"^\d", raw_rows[0][0] if raw_rows[0] else ""):
                roles = {"num": 0, "name": 1, "req": 2, "meth": 3 if len(raw_rows[0]) > 3 else 2}
                data_start = 0
            else:
                continue

        name_i = roles.get("name", 1)
        num_i = roles.get("num", 0)
        req_i = roles.get("req")
        meth_i = roles.get("meth")
        applies_i = roles.get("applies")

        for row in raw_rows[data_start:]:
            if not row or _is_header_row(row) or _is_section_row(row):
                continue
            if max(name_i, num_i) >= len(row):
                continue
            num_cell = row[num_i] if num_i < len(row) else ""
            name_cell = row[name_i] if name_i < len(row) else ""
            req_cell = row[req_i] if req_i is not None and req_i < len(row) else ""
            meth_cell = row[meth_i] if meth_i is not None and meth_i < len(row) else ""
            applies = row[applies_i] if applies_i is not None and applies_i < len(row) else ""

            if not name_cell:
                continue
            # num may be multi: «2.1.1, 2.1.2»
            mnum = re.match(r"^(\d+)", num_cell.strip())
            if mnum:
                order = int(mnum.group(1))
            else:
                m2 = _RE_NUM_START.match(name_cell.strip())
                if m2:
                    order = int(m2.group(1))
                    name_cell = m2.group(2)
                else:
                    auto_order += 1
                    order = auto_order

            name_cell = _normalize_ws(name_cell)
            if len(name_cell) < 3:
                continue

            req_doc, req_clause = _split_doc_and_clause(req_cell, default_tu)
            meth_doc, meth_clause = _split_doc_and_clause(meth_cell, default_tu)
            if meth_clause and not meth_doc and default_tu:
                meth_doc = default_tu

            names = _split_jammed_test_name(name_cell)
            req_toks = _clause_tokens(req_clause)
            meth_toks = _clause_tokens(meth_clause)

            if len(names) == 2 and len(req_toks) >= 2 and len(meth_toks) >= 2:
                # pair each name with one clause
                for j, nm in enumerate(names):
                    items.append(
                        ParsedProgramItem(
                            sort_order=order if j == 0 else order,
                            name=nm,
                            requirement_doc=req_doc,
                            requirement_clause=req_toks[j] if j < len(req_toks) else req_clause,
                            method_doc=meth_doc,
                            method_clause=meth_toks[j] if j < len(meth_toks) else meth_clause,
                            applies_to=applies,
                        )
                    )
            else:
                items.append(
                    ParsedProgramItem(
                        sort_order=order,
                        name=name_cell,
                        requirement_doc=req_doc,
                        requirement_clause=req_clause,
                        method_doc=meth_doc,
                        method_clause=meth_clause,
                        applies_to=applies,
                    )
                )

    # dedupe
    seen: set[tuple[int, str]] = set()
    out: list[ParsedProgramItem] = []
    for it in sorted(items, key=lambda x: (x.sort_order, x.name)):
        key = (it.sort_order, it.name.lower())
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


def _parse_text_body_items(paragraphs: list[str], default_tu: str) -> list[ParsedProgramItem]:
    """Испытания, описанные только текстом (п. 2.4.1 …)."""
    items: list[ParsedProgramItem] = []
    order = 0
    for p in paragraphs:
        m = _RE_TEXT_ITEM.match(p)
        if not m:
            continue
        clause, title = m.group(1), _normalize_ws(m.group(2))
        # skip if title is too generic continuation
        if len(title) < 8:
            continue
        order += 1
        # extract TU refs in line
        req_doc, _ = _split_doc_and_clause(p, default_tu)
        items.append(
            ParsedProgramItem(
                sort_order=order,
                name=title.rstrip("."),
                requirement_doc=req_doc or default_tu,
                requirement_clause=clause,
                method_doc="",
                method_clause="",
            )
        )
    return items


def parse_program_docx(path: Path | str) -> ParsedProgram:
    """Полный разбор одного .docx программы испытаний."""
    from docx import Document

    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    if file_path.suffix.lower() != ".docx":
        raise ValueError("Поддерживается только .docx (не .doc / PDF)")

    doc = Document(str(file_path))
    paragraphs = [
        _normalize_ws(p.text) for p in doc.paragraphs if p.text and p.text.strip()
    ]
    header = _parse_header(paragraphs)
    default_tu = ""
    if header["tu_ref"]:
        default_tu = header["tu_ref"].split(",")[0].strip()

    table_items = _parse_table_items(doc, default_tu)
    text_items: list[ParsedProgramItem] = []
    if len(table_items) < 3:
        # мало из таблиц — дополняем/берём из текста
        text_items = _parse_text_body_items(paragraphs, default_tu)

    if table_items and text_items:
        # prefer tables; add text items with names not already present
        existing = {it.name.lower() for it in table_items}
        for it in text_items:
            if it.name.lower() not in existing:
                table_items.append(it)
        items = table_items
    elif table_items:
        items = table_items
    else:
        items = text_items

    notes = []
    notes.append(f"марок: {len(header['cable_marks'])}")
    notes.append(f"пунктов: {len(items)}")
    notes.append(f"абзацев: {len(paragraphs)}, таблиц: {len(doc.tables)}")
    if not header["cable_marks"]:
        notes.append("⚠ марка не извлечена — проверьте шапку/список «1) МАРКА – …»")
    if not items:
        notes.append("⚠ испытания не найдены")

    return ParsedProgram(
        name=header["name"] or file_path.stem,
        test_type=header["test_type"],
        cable_mark_text=header["cable_mark_text"],
        cable_marks=list(header["cable_marks"]),
        tu_ref=header["tu_ref"],
        source_path=str(file_path.resolve()),
        notes="; ".join(notes),
        items=items,
    )


def import_program_from_docx(
    path: Path | str,
    *,
    db_path: Path | str | None = None,
    match_price: bool = True,
) -> dict[str, Any]:
    """Парсит один DOCX и сохраняет в БД."""
    from ..persistence.sqlite_repo import (
        DB_PATH_DEFAULT,
        create_test_program,
        match_program_items_to_price,
    )

    db = db_path or DB_PATH_DEFAULT
    parsed = parse_program_docx(path)
    items_payload = [
        {
            "sort_order": it.sort_order,
            "name": it.name,
            "requirement_doc": it.requirement_doc,
            "requirement_clause": it.requirement_clause,
            "method_doc": it.method_doc,
            "method_clause": it.method_clause,
            "meta": {"applies_to": it.applies_to} if it.applies_to else None,
        }
        for it in parsed.items
    ]
    program_id = create_test_program(
        name=parsed.name,
        test_type=parsed.test_type,
        cable_mark_text=parsed.cable_mark_text,
        tu_ref=parsed.tu_ref,
        source_path=parsed.source_path,
        notes=parsed.notes,
        items=items_payload,
        db_path=db,
    )
    match_stats: dict[str, Any] = {
        "matched": 0,
        "unmatched": 0,
        "total": 0,
        "rate": 0.0,
        "summary": "сопоставлено 0/0",
    }
    if match_price and items_payload:
        match_stats = match_program_items_to_price(program_id, db_path=db)
    return {
        "program_id": program_id,
        "name": parsed.name,
        "test_type": parsed.test_type,
        "cable_mark_text": parsed.cable_mark_text,
        "cable_marks": parsed.cable_marks,
        "tu_ref": parsed.tu_ref,
        "items_count": len(parsed.items),
        "matched": match_stats.get("matched", 0),
        "unmatched": match_stats.get("unmatched", 0),
        "total": match_stats.get("total", 0),
        "rate": match_stats.get("rate", 0.0),
        "summary": match_stats.get("summary", ""),
        "source_path": parsed.source_path,
        "notes": parsed.notes,
    }
