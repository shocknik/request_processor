"""
pdf_extractor.py — извлечение данных из входящих заявок (PDF, Word .docx).

Текстовые PDF: pdfplumber.
Сканы (картинки без текстового слоя): OCR через pytesseract или easyocr.
Поиск марок — по структуре строки (бренд + «NхM»), без списка брендов.
Организации — через organization_extractor (заказчик, производитель).
"""

from __future__ import annotations

import logging
import re
import shutil
from typing import Literal
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

from .cable_mark_parser import extract_document_from_text, fix_ocr_document_text
from .models import CableMarkMatch, PdfExtractionResult
from .organization_extractor import (
    extract_organizations,
    pick_customer_name,
    pick_manufacturer_name,
)

logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None  # type: ignore[assignment]

# --- Паттерны марки (структурные, без whitelist брендов) ---

_SIZE_PART = (
    r"\d+\s*[зЗпП]?\s*[хx]\s*"
    r"(?:[\d.,\(\)]+|[а-яёa-zA-Z]{1,6})"
    r"(?:\s*[хx]\s*[\d.,\(\)]+)*"
    r"(?:[а-яёa-zA-Z\-\(\),\d/]*)"
    r"(?:\s*\([^)]+\))?"
    r"(?:-\d+[.,]?\d*)?"
)

_NAME_PART = (
    r"(?:ККЗ\s+МК\s+)?"
    r"[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z0-9\-\(\)/]+"
    r"(?:[\-–/][А-ЯЁа-яёA-Za-z0-9\(\)/]+)*"
)

# Размер: 2х2, 2x2x0,52, 4x2x0.52 (LAN-кабель)
_SIZE_PART_LATIN = (
    r"\d+\s*x\s*\d+"
    r"(?:\s*x\s*[\d.,]+)?"
)

# СПЕЦЛАН F/UTP … 2x2x0,52
_SPECLAN_MARK_PATTERN = re.compile(
    r"(?:\d+\.\s*)?"
    r"(СПЕЦЛАН\s+(?:SF?/)?UTP\s+Cat\s+5\w\s+ZH\s+нг\(А\)-HF\s+\d+\s*x\s*\d+(?:\s*x\s*[\d.,]+)?)",
    re.IGNORECASE,
)

# Нумерованный список марок в письме
_LETTER_MARKS_BLOCK = re.compile(
    r"марк[аи]\s+кабел[ья][:\s]+(.+?)(?=Последующ|С\s+уважением|Суважением|$)",
    re.IGNORECASE | re.DOTALL,
)
_LETTER_MARK_ITEM = re.compile(
    r"\d+\.\s*"
    r"(СПЕЦЛАН\s+.+?)"
    r"(?:\s+(ТУ\s*\d+\.[КкKk]\d{2,3}-\d{3}-\d{4}))?"
    r"(?=\s+В\s+количестве|;|\d+\.\s+СПЕЦ|\d+\.\s+[А-ЯЁA-Z]|$)",
    re.IGNORECASE,
)
_TU_TAIL_PATTERN = re.compile(
    r"ТУ\s*\d+\.[КкKk]\d{2,3}-\d{3}-\d{4}",
    re.IGNORECASE,
)

_MARK_PATTERN = re.compile(
    rf"{_NAME_PART}\s+{_SIZE_PART}",
    re.IGNORECASE,
)

# Контекстный поиск после «марки:» / «марка»
_AFTER_MARKI_PATTERN = re.compile(
    rf"(?:кабел[ья]\s+)?(?:силовой\s+)?марк[аи]\s*:?\s*"
    rf"({_NAME_PART}\s+{_SIZE_PART})",
    re.IGNORECASE,
)

# Провод/кабель «марки X» (включая OCR «ПровоА»)
_PRODUCT_MARK_PATTERN = re.compile(
    rf"(?:Прово\w*|Кабел\w*)[^\n]{{0,50}}?марк[аи]\s*:?\s*"
    rf"({_NAME_PART}\s+{_SIZE_PART})",
    re.IGNORECASE,
)

_REJECT_PREFIXES = re.compile(
    r"^(?:солнечного|излучения|воздействию|стойкость|требования|наименование|"
    r"марка|образец|№|п\.|пункт|гост|ту|нд)\b",
    re.IGNORECASE,
)

DEFAULT_OCR_DPI = 200


def _require_pdfplumber() -> None:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber не установлен. Выполни: pip install pdfplumber")


def _normalize_text(text: str) -> str:
    """Приводит текст PDF/OCR к удобному для поиска виду."""
    text = text.replace("\xa0", " ")
    text = text.replace("—", "-").replace("–", "-")
    text = re.sub(r"-\s*\n\s*", "-", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _normalize_text_for_marks(text: str) -> str:
    """
    Нормализация для поиска марок.

    Сохраняет латинское «x» в размерах (2x2x0,52, Cat 5e) и в F/UTP.
    Кириллическое «х» унифицирует в «x» для единого паттерна.
    """
    text = _normalize_text(text)
    text = text.replace("×", "x").replace("Х", "x").replace("х", "x")
    return text


def _fix_ocr_confusables(text: str) -> str:
    """Исправляет типичные OCR-ошибки перед поиском марок и ТУ."""
    text = fix_ocr_document_text(text)
    text = re.sub(r"(?<=[А-ЯЁA-Zа-яё])l(?=[хx×]|\d)", "1", text)
    text = re.sub(r"(?<=[А-ЯЁA-Zа-яё])I(?=[хx×]|\d)", "1", text)
    text = re.sub(r"\bl(?=[хx×]\s*[\d.,])", "1", text, flags=re.IGNORECASE)
    text = re.sub(r"\bI(?=[хx×]\s*[\d.,])", "1", text)
    text = re.sub(r"\bЗ\s*х\b", "3х", text, flags=re.IGNORECASE)
    text = re.sub(r"(\d+)\s*х\s*ок\b", r"\1х4ок", text, flags=re.IGNORECASE)
    text = re.sub(r"N;\s*PE", "(N,PE)", text, flags=re.IGNORECASE)
    text = re.sub(r"\)J\b", ")", text)
    text = re.sub(r"PEJ", "PE)", text, flags=re.IGNORECASE)
    return text


def _is_plausible_mark(mark: str) -> bool:
    """Отсекает явные ложные срабатывания из «рваного» текста PDF/OCR."""
    if re.search(r"\d{2}\.\d{4}", mark):
        return False
    if _REJECT_PREFIXES.search(mark):
        return False
    if re.match(r"^нг\(", mark, re.IGNORECASE):
        return False
    if not re.match(r"^(?:ККЗ\s+МК\s+|СПЕЦЛАН\s+)?[А-ЯЁA-Z]", mark, re.IGNORECASE):
        return False
    has_cyr_size = re.search(
        r"\d+\s*[зЗпП]?\s*[хx]\s*(?:[\d.,]|[а-яёa-zA-Z])", mark, re.IGNORECASE
    )
    has_lan_size = re.search(_SIZE_PART_LATIN, mark, re.IGNORECASE)
    if not has_cyr_size and not has_lan_size:
        return False
    name = re.split(r"\s+\d", mark, maxsplit=1)[0]
    if len(name) > 100 or len(name) < 2:
        return False
    return True


def _clean_mark(raw: str) -> str:
    """Убирает хвостовой мусор из кандидата в марку."""
    mark = raw.strip(" .,;:\n")
    mark = re.sub(
        rf"^(.+?{_SIZE_PART})\s+м\s+\d.*",
        r"\1",
        mark,
        flags=re.IGNORECASE,
    )
    mark = re.sub(r"\s+Упаковка.*$", "", mark, flags=re.IGNORECASE)
    mark = re.sub(
        r"\s+(?:Прово\w*|Кабел\w*|марк[аи]|ТУ|ГОСТ|СТО).*$",
        "",
        mark,
        flags=re.IGNORECASE,
    )
    if mark.upper().startswith("СПЕЦЛАН"):
        mark = re.sub(r"\s+В\s+количестве.*$", "", mark, flags=re.IGNORECASE)
    return mark.strip()


def _context_snippet(text: str, start: int, end: int, radius: int = 120) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    snippet = text[left:right].strip()
    return re.sub(r"\s+", " ", snippet)


def _add_match(
    matches: list[CableMarkMatch],
    seen: set[str],
    mark: str,
    text: str,
    start: int,
    end: int,
    *,
    document: str | None = None,
) -> None:
    mark = _clean_mark(mark)
    key = mark.lower()
    if len(mark) < 5 or key in seen or not _is_plausible_mark(mark):
        return
    seen.add(key)
    context = _context_snippet(text, start, end, radius=180)
    doc = document or extract_document_from_text(context)
    matches.append(
        CableMarkMatch(
            mark=mark,
            context=context,
            document=doc,
        )
    )


def _find_letter_list_marks(text: str) -> list[tuple[str, int, int, str | None]]:
    """Марки из нумерованного списка в гарантийном/сопроводительном письме."""
    found: list[tuple[str, int, int, str | None]] = []
    block = _LETTER_MARKS_BLOCK.search(text)
    search_in = block.group(1) if block else text
    base_offset = block.start(1) if block else 0

    for m in _LETTER_MARK_ITEM.finditer(search_in):
        mark = m.group(1).strip()
        if not mark.upper().startswith("СПЕЦЛАН"):
            continue
        doc = m.group(2).strip() if m.lastindex and m.lastindex >= 2 and m.group(2) else None
        if not doc:
            tail = search_in[m.end() : m.end() + 80]
            tu_m = _TU_TAIL_PATTERN.search(tail)
            doc = tu_m.group(0) if tu_m else None
        if doc:
            doc = extract_document_from_text(doc) or doc
        found.append((mark, base_offset + m.start(1), base_offset + m.end(1), doc))

    if not found:
        for m in _SPECLAN_MARK_PATTERN.finditer(text):
            tail = text[m.end() : m.end() + 80]
            tu_m = _TU_TAIL_PATTERN.search(tail)
            doc = extract_document_from_text(tu_m.group(0)) if tu_m else None
            found.append((m.group(1), m.start(1), m.end(1), doc))

    return found


def find_cable_marks(text: str) -> list[CableMarkMatch]:
    """
    Ищет марки кабелей по структурному паттерну «название + NхM».

    Не использует список брендов — любая строка нужной формы.
    Поддерживает LAN (СПЕЦЛАН F/UTP 2x2x0,52) и нумерованные списки в письмах.
    """
    normalized = _fix_ocr_confusables(_normalize_text_for_marks(text))
    if not normalized:
        return []

    seen: set[str] = set()
    matches: list[CableMarkMatch] = []

    for mark, start, end, doc in _find_letter_list_marks(normalized):
        _add_match(matches, seen, mark, normalized, start, end, document=doc)

    for pattern in (_SPECLAN_MARK_PATTERN, _AFTER_MARKI_PATTERN, _PRODUCT_MARK_PATTERN, _MARK_PATTERN):
        for m in pattern.finditer(normalized):
            raw = m.group(1) if m.lastindex else m.group(0)
            _add_match(matches, seen, raw, normalized, m.start(), m.end())

    return matches


def _find_tesseract() -> str | None:
    """Ищет tesseract.exe в PATH и типичных путях Windows."""
    found = shutil.which("tesseract")
    if found:
        return found
    project_root = Path(__file__).resolve().parents[2]
    for candidate in (
        project_root / "tools" / "Tesseract-OCR" / "tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ):
        path = Path(candidate)
        if path.exists():
            return str(path)
    return None


def _render_pages(pdf_path: Path, dpi: int = DEFAULT_OCR_DPI) -> list:
    """Рендерит страницы PDF в PIL.Image через PyMuPDF."""
    import fitz
    from PIL import Image

    doc = fitz.open(str(pdf_path))
    images: list = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    for page in doc:
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        images.append(Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples))
    doc.close()
    return images


def _ocr_image_tesseract(image, pytesseract) -> str:
    return pytesseract.image_to_string(image, lang="rus+eng")


def _ocr_with_tesseract(pdf_path: Path, dpi: int = DEFAULT_OCR_DPI) -> str:
    import pytesseract

    tesseract_cmd = _find_tesseract()
    if not tesseract_cmd:
        raise RuntimeError(
            "Tesseract OCR не найден. Установи: "
            "https://github.com/UB-Mannheim/tesseract/wiki "
            "(нужен язык rus)."
        )
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    images = _render_pages(pdf_path, dpi=dpi)
    if len(images) == 1:
        text = _ocr_image_tesseract(images[0], pytesseract)
        return text.strip()

    parts: list[str] = [""] * len(images)
    with ThreadPoolExecutor(max_workers=min(4, len(images))) as pool:
        futures = {
            pool.submit(_ocr_image_tesseract, img, pytesseract): idx
            for idx, img in enumerate(images)
        }
        for future in as_completed(futures):
            idx = futures[future]
            parts[idx] = future.result().strip()

    return "\n\n".join(p for p in parts if p)


def _ocr_with_easyocr(pdf_path: Path, dpi: int = DEFAULT_OCR_DPI) -> str:
    import easyocr
    import numpy as np

    reader = easyocr.Reader(["ru", "en"], gpu=False, verbose=False)
    parts: list[str] = []
    for image in _render_pages(pdf_path, dpi=dpi):
        lines = reader.readtext(np.array(image), detail=0, paragraph=True)
        if lines:
            parts.append("\n".join(lines))
    return "\n\n".join(parts)


def ocr_pdf(pdf_path: Path | str, dpi: int = DEFAULT_OCR_DPI) -> str:
    """
    Распознаёт текст со скана. Сначала tesseract, при недоступности — easyocr.
    """
    path = Path(pdf_path)
    if _find_tesseract():
        try:
            return _ocr_with_tesseract(path, dpi=dpi)
        except Exception as exc:
            logger.warning("Tesseract OCR failed: %s", exc)

    try:
        return _ocr_with_easyocr(path, dpi=dpi)
    except ImportError as exc:
        raise RuntimeError(
            "Для сканов нужен OCR. Установи Tesseract (rus) или: pip install easyocr"
        ) from exc


def extract_text(pdf_path: Path | str, *, use_ocr: bool = False) -> str:
    """Извлекает текст: pdfplumber, для сканов — с OCR если use_ocr=True."""
    _require_pdfplumber()
    path = Path(pdf_path)
    parts: list[str] = []

    with pdfplumber.open(path) as doc:
        for page in doc.pages:
            page_text = page.extract_text() or ""
            if page_text:
                parts.append(page_text)

    text = "\n\n".join(parts)
    if not text.strip() and use_ocr:
        text = ocr_pdf(path)
    return text


def extract_tables(pdf_path: Path | str) -> list[list[list[str]]]:
    """Извлекает таблицы. Каждая ячейка — строка (пустые → '')."""
    _require_pdfplumber()
    path = Path(pdf_path)
    tables: list[list[list[str]]] = []

    with pdfplumber.open(path) as doc:
        for page in doc.pages:
            for table in page.extract_tables() or []:
                cleaned = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in table
                ]
                if any(any(cell for cell in row) for row in cleaned):
                    tables.append(cleaned)

    return tables


def _tables_to_text(tables: list[list[list[str]]]) -> str:
    return "\n".join(" ".join(cell for cell in row if cell) for table in tables for row in table)


def _detect_scanned(pdf_path: Path) -> tuple[bool, int]:
    """True, если PDF похож на скан (есть изображения, но нет текста)."""
    _require_pdfplumber()
    with pdfplumber.open(pdf_path) as doc:
        page_count = len(doc.pages)
        has_images = any(page.images for page in doc.pages)
        text_len = sum(len(page.extract_text() or "") for page in doc.pages)
        is_scanned = page_count > 0 and text_len == 0 and has_images
        return is_scanned, page_count


def extract_from_pdf(
    pdf_path: Path | str,
    *,
    use_ocr: bool = True,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> PdfExtractionResult:
    """
    Полное извлечение: текст, таблицы, марки кабелей, метаданные.

    Для сканов при use_ocr=True автоматически запускает OCR.
    """
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF не найден: {path}")

    is_scanned, page_count = _detect_scanned(path)
    ocr_used = False
    tables: list[list[list[str]]] = []

    if is_scanned and use_ocr:
        text = ocr_pdf(path, dpi=ocr_dpi)
        ocr_used = True
    else:
        text = extract_text(path)
        tables = extract_tables(path)

    search_text = text
    if tables:
        search_text = f"{text}\n{_tables_to_text(tables)}".strip()

    cable_marks = find_cable_marks(search_text)
    organizations = extract_organizations(search_text)

    logger.info(
        "PDF %s: pages=%d, text=%d chars, tables=%d, marks=%d, orgs=%d, scanned=%s, ocr=%s",
        path.name,
        page_count,
        len(text),
        len(tables),
        len(cable_marks),
        len(organizations),
        is_scanned,
        ocr_used,
    )

    return PdfExtractionResult(
        source_path=str(path.resolve()),
        source_type="pdf",
        page_count=page_count,
        text=text,
        tables=tables,
        cable_marks=cable_marks,
        organizations=organizations,
        customer_name=pick_customer_name(organizations),
        manufacturer_name=pick_manufacturer_name(organizations),
        is_scanned=is_scanned,
        ocr_used=ocr_used,
        extracted_at=datetime.now(),
    )


def extract_text_from_docx(docx_path: Path | str) -> str:
    """Извлекает текст из Word-документа (.docx)."""
    from docx import Document

    path = Path(docx_path)
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


def _build_extraction_result(
    *,
    source_path: Path,
    source_type: Literal["pdf", "docx"],
    text: str,
    page_count: int = 0,
    tables: list[list[list[str]]] | None = None,
    is_scanned: bool = False,
    ocr_used: bool = False,
) -> PdfExtractionResult:
    """Собирает результат: марки + организации из текста заявки."""
    search_text = text
    if tables:
        search_text = f"{text}\n{_tables_to_text(tables)}".strip()
    organizations = extract_organizations(search_text)
    return PdfExtractionResult(
        source_path=str(source_path.resolve()),
        source_type=source_type,
        page_count=page_count,
        text=text,
        tables=tables or [],
        cable_marks=find_cable_marks(search_text),
        organizations=organizations,
        customer_name=pick_customer_name(organizations),
        manufacturer_name=pick_manufacturer_name(organizations),
        is_scanned=is_scanned,
        ocr_used=ocr_used,
        extracted_at=datetime.now(),
    )


def extract_from_document(
    path: Path | str,
    *,
    use_ocr: bool = True,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> PdfExtractionResult:
    """
    Единая точка входа для заявок: PDF или Word (.docx).

    PDF делегирует в extract_from_pdf; Word — текст + марки + организации.
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_from_pdf(file_path, use_ocr=use_ocr, ocr_dpi=ocr_dpi)
    if suffix == ".docx":
        text = extract_text_from_docx(file_path)
        return _build_extraction_result(
            source_path=file_path,
            source_type="docx",
            text=text,
        )
    raise ValueError(f"Неподдерживаемый формат: {suffix}. Используйте PDF или .docx")