"""
Макет протокола испытаний по шаблону Word.

Заполняет известные поля (заказчик, изготовитель, объект/марки, цель).
Это черновик для оператора, не финальный протокол с результатами.
"""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

from ..config import GENERATED_DIR, PROTOCOL_TEMPLATE_NAME, TEMPLATES_DIR
from ..persistence.sqlite_repo import DB_PATH_DEFAULT, get_order_details

TEMPLATE_DEFAULT = TEMPLATES_DIR / PROTOCOL_TEMPLATE_NAME


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """Простая подстановка по полному тексту параграфа (с сохранением runs грубо)."""
    text = paragraph.text
    if not text:
        return
    new = text
    for old, val in mapping.items():
        if old in new and val:
            new = new.replace(old, val)
    if new != text:
        # очищаем runs и пишем одним
        if paragraph.runs:
            paragraph.runs[0].text = new
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new)


def _set_paragraph_if_label(paragraph, label_prefixes: tuple[str, ...], value: str) -> bool:
    """Если параграф — метка вида «юридический адрес:», дописывает значение."""
    raw = paragraph.text.strip()
    if not raw or not value:
        return False
    lower = raw.lower()
    for pref in label_prefixes:
        if lower.startswith(pref.lower()):
            # уже заполнено чем-то кроме двоеточия
            if len(raw) > len(pref) + 2 and not raw.endswith(":"):
                return False
            new_text = f"{pref.rstrip(':')}: {value}"
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(new_text)
            return True
    return False


def _marks_summary(marks: list[dict]) -> str:
    names = [str(m.get("mark") or "").strip() for m in marks if m.get("mark")]
    names = [n for n in names if n]
    if not names:
        return "—"
    if len(names) == 1:
        return names[0]
    if len(names) <= 3:
        return ", ".join(names)
    return ", ".join(names[:3]) + f" и ещё {len(names) - 3}"


def generate_protocol_draft_from_order(
    order_id: int,
    output_path: Path | str | None = None,
    *,
    template_path: Path | str = TEMPLATE_DEFAULT,
    db_path: Path | str = DB_PATH_DEFAULT,
) -> Path:
    """
    Копирует шаблон протокола и подставляет реквизиты заказа.
    Возвращает путь к .docx-черновику.
    """
    details = get_order_details(order_id, db_path=db_path)
    if not details:
        raise ValueError(f"Заказ №{order_id} не найден")

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Шаблон протокола не найден: {template}")

    customer = (details.get("customer_name") or "заказчик").strip()
    manufacturer = (details.get("manufacturer_name") or customer).strip()
    subject = (details.get("subject") or "проведение испытаний").strip()
    marks = details.get("marks") or []
    marks_text = _marks_summary(marks)
    customer_addr = (
        details.get("customer_legal_address")
        or details.get("customer_address")
        or details.get("customer_actual_address")
        or ""
    )
    mfg_addr = (
        details.get("manufacturer_legal_address")
        or details.get("manufacturer_address")
        or details.get("manufacturer_actual_address")
        or ""
    )

    if output_path is None:
        from .document_pack import safe_filename_part

        safe = safe_filename_part(customer, max_len=30, default="заказчик")
        out_dir = GENERATED_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / (
            f"Протокол_макет_{safe}_заказ{order_id}_{datetime.now():%Y%m%d_%H%M}.docx"
        )
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    shutil.copy2(template, output_path)
    doc = Document(str(output_path))

    # Параграфы с известными формулировками
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue

        if t.startswith("Наименование объекта испытаний"):
            continue

        if "На испытания представлен" in t or "На испытания представлен(-ы)" in t:
            new = (
                f"На испытания представлен(-ы) образец(-цы) кабеля {marks_text}."
            )
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
            continue

        if t.startswith("Образец(-цы) изготовлен") or "изготовлен(-ы)" in t and "по ТУ" in t:
            docs = sorted(
                {
                    str(m.get("document") or "").strip()
                    for m in marks
                    if m.get("document")
                }
            )
            tu = docs[0] if docs else "ТУ …"
            new = (
                f"Образец(-цы) изготовлен(-ы) {manufacturer} по {tu}."
            )
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
            continue

        if t.startswith("Определение соответствия") or t.startswith("Определение (наименование"):
            new = (
                f"Определение соответствия образца(-ов) кабеля {marks_text} "
                f"требованиям НД / цели: {subject}."
            )
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
            continue

        # Блок заказчика: строка «Наименование (уникальный номер…» сразу после заголовка
        if t.startswith("Наименование (уникальный номер") and customer:
            new = f"Наименование: {customer}"
            if para.runs:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""
            continue

        if t == "Наименование" and manufacturer:
            if para.runs:
                para.runs[0].text = f"Наименование: {manufacturer}"
                for r in para.runs[1:]:
                    r.text = ""
            continue

        _set_paragraph_if_label(
            para,
            ("юридический адрес:", "юридический адрес"),
            customer_addr or mfg_addr,
        )

    # Таблица шапки: номер протокола-заготовки
    if doc.tables:
        header = doc.tables[0]
        for row in header.rows:
            for cell in row.cells:
                if "ПРОТОКОЛ" in cell.text.upper():
                    # не затираем весь merged-блок полностью — помечаем заказ
                    for p in cell.paragraphs:
                        if "ПРОТОКОЛ" in p.text.upper():
                            note = f"ПРОТОКОЛ (макет) · заказ №{order_id} · {datetime.now():%d.%m.%Y}"
                            if p.runs:
                                p.runs[0].text = note
                                for r in p.runs[1:]:
                                    r.text = ""
                            break

    doc.save(str(output_path))
    return output_path
