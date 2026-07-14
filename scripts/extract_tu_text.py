"""
Извлечение текста ТУ → data/knowledge/manufacturer_v1/raw_text/<safe_id>.txt

- .docx: python-docx
- .doc: Word COM (win32com), если доступен
- .pdf: pdfplumber (без OCR)

Запуск:
  .venv\\Scripts\\python.exe scripts\\extract_tu_text.py
  .venv\\Scripts\\python.exe scripts\\extract_tu_text.py --only-docx
  .venv\\Scripts\\python.exe scripts\\extract_tu_text.py --limit 10
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

TU_DIR = ROOT / "data" / "training" / "rag_corpus" / "tu"
OUT_DIR = ROOT / "data" / "knowledge" / "manufacturer_v1" / "raw_text"
INDEX_PATH = ROOT / "data" / "knowledge" / "manufacturer_v1" / "tu_index.yaml"


def safe_id(tu_id: str | None, file_name: str) -> str:
    base = tu_id or Path(file_name).stem
    base = re.sub(r'[<>:"/\\|?*\s]+', "_", base)
    return base[:80] or "unknown"


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                # unique preserve order (merged cells duplicate)
                seen: set[str] = set()
                uniq: list[str] = []
                for c in cells:
                    if c not in seen:
                        seen.add(c)
                        uniq.append(c)
                parts.append(" | ".join(uniq))
    return "\n".join(parts)


def extract_doc_com(path: Path) -> str:
    """Legacy .doc via Word COM."""
    try:
        import win32com.client  # type: ignore
    except ImportError as exc:
        raise RuntimeError("pywin32 not installed") from exc

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        try:
            text = doc.Content.Text
        finally:
            doc.Close(False)
    finally:
        word.Quit()
    return (text or "").replace("\r", "\n")


def extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t.strip())
    return "\n\n".join(parts)


def extract_one(path: Path) -> tuple[str, str]:
    """Returns (text, method)."""
    suf = path.suffix.lower()
    if suf == ".docx":
        return extract_docx(path), "docx"
    if suf == ".doc":
        return extract_doc_com(path), "doc_com"
    if suf == ".pdf":
        return extract_pdf(path), "pdfplumber"
    if suf == ".rtf":
        # bare strip — better than nothing
        raw = path.read_bytes()
        try:
            text = raw.decode("utf-8", errors="ignore")
        except Exception:
            text = raw.decode("cp1251", errors="ignore")
        text = re.sub(r"\\[a-z]+\d* ?", " ", text)
        text = re.sub(r"[{}]", " ", text)
        return re.sub(r"\s+", " ", text).strip(), "rtf_naive"
    raise ValueError(f"unsupported: {suf}")


def load_index() -> dict:
    if not INDEX_PATH.is_file():
        return {"documents": []}
    import yaml

    return yaml.safe_load(INDEX_PATH.read_text(encoding="utf-8")) or {"documents": []}


def save_index(data: dict) -> None:
    import yaml

    INDEX_PATH.write_text(
        yaml.safe_dump(data, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--only-docx", action="store_true")
    ap.add_argument("--only-doc", action="store_true")
    ap.add_argument("--limit", type=int, default=0)
    ap.add_argument("--force", action="store_true", help="перезаписать raw_text")
    args = ap.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    index = load_index()
    by_name = {d["file_name"]: d for d in index.get("documents") or []}

    files = sorted(TU_DIR.iterdir(), key=lambda p: p.name.lower())
    files = [p for p in files if p.is_file()]
    if args.only_docx:
        files = [p for p in files if p.suffix.lower() == ".docx"]
    if args.only_doc:
        files = [p for p in files if p.suffix.lower() == ".doc"]
    if args.limit:
        files = files[: args.limit]

    ok = fail = skip = 0
    for path in files:
        meta = by_name.get(path.name) or {
            "file_name": path.name,
            "tu_id": None,
            "ext": path.suffix.lower(),
            "status": "pending",
            "brands_hint": [],
            "notes": "",
        }
        sid = safe_id(meta.get("tu_id"), path.name)
        out = OUT_DIR / f"{sid}.txt"
        if out.exists() and not args.force:
            skip += 1
            if meta.get("status") == "pending":
                meta["status"] = "inventoried"
                meta["raw_text"] = str(out.relative_to(ROOT)).replace("\\", "/")
                by_name[path.name] = meta
            continue
        try:
            text, method = extract_one(path)
            if len(text.strip()) < 40:
                raise RuntimeError(f"too short text ({len(text)} chars) via {method}")
            out.write_text(text, encoding="utf-8")
            meta["status"] = "inventoried"
            meta["raw_text"] = str(out.relative_to(ROOT)).replace("\\", "/")
            meta["extract_method"] = method
            meta["text_chars"] = len(text)
            by_name[path.name] = meta
            ok += 1
            print(f"OK  {path.name[:60]}  ({method}, {len(text)} chars)")
        except Exception as exc:
            fail += 1
            meta["notes"] = f"extract_failed: {exc}"
            by_name[path.name] = meta
            print(f"FAIL {path.name[:60]}  {exc}")

    index["documents"] = list(by_name.values())
    index["count"] = len(index["documents"])
    save_index(index)
    print(f"\nDone: ok={ok} skip={skip} fail={fail} → {OUT_DIR}")


if __name__ == "__main__":
    main()
