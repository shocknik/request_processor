"""
pdf_extractor.py — извлечение данных из входящих заявок (PDF, Word .docx).

Текстовые PDF: pdfplumber.
Сканы (картинки без текстового слоя): OCR через pytesseract или easyocr.
Поиск марок — по структуре строки (бренд + «NхM»), без списка брендов.
Организации — через organization_extractor (заказчик, производитель).
"""

from __future__ import annotations

import hashlib
import logging
import re
import shutil
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Literal

from ..config import OCR_CACHE_DIR
from ..models import CableMarkMatch, PdfExtractionResult
from ..parsing.cable_mark_parser import extract_document_from_text, fix_ocr_document_text
from .ocr_mark_normalizer import normalize_mark_after_ocr
from .ocr_text_normalizer import normalize_ocr_text
from .organization_extractor import (
    extract_organizations,
    finalize_organizations,
    pick_customer_name,
    pick_manufacturer_name,
)

logger = logging.getLogger(__name__)

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None  # type: ignore[assignment]

# --- Паттерны марки (структурные, без whitelist брендов) ---

_SIZE_PART = (
    r"\d+\s*[зЗпП]?\s*[хx]\s*"
    r"(?:[\d.,\(\)]+|[а-яёa-zA-Z]{1,6})"
    r"(?:\s*[хx]\s*[\d.,\(\)]+)*"
    # Без пробелов: «(N, PE)» нормализуется в «(N,PE)» в _fix_periodic_letter_ocr
    r"(?:[а-яёa-zA-Z\-\(\),\d/]*)"
    r"(?:\s*\([^)]+\))?"
    r"(?:-\d+[.,]?\d*)?"
)

_NAME_PART = (
    r"(?:ККЗ\s+МК\s+)?"
    r"[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z0-9\-\(\)/]+"
    r"(?:[\-–/][А-ЯЁа-яёA-Za-z0-9\(\)/]+)*"
)

# Размер: 2х2, 2x2x0,52, 4x2x0.52 (LAN-кабель)
_SIZE_PART_LATIN = (
    r"\d+\s*x\s*\d+"
    r"(?:\s*x\s*[\d.,]+)?"
)

# СПЕЦЛАН / SPECLAN / CMELVIAH (OCR) F/UTP … 2x2x0,52
_LAN_BRAND = r"(?:СПЕЦЛАН|SPECLAN|CMELVIAH|CMELAN|СNЕ[UW][АA]{1,2}Н|СNЕLWIАН)"
_LAN_FIRE = r"(?:нг|ur|Hr|hr|ng|Нг)\s*\(\s*A\s*\)"
_LAN_MARK_PATTERN = re.compile(
    r"(?:\d+\.\s*)?"
    rf"({_LAN_BRAND}\s+(?:SF?/)?UTP\s+Cat\s+5\w\s+ZH\s+{_LAN_FIRE}-HF\s+\d+\s*x\s*\d+(?:\s*x\s*[\d.,]+)?)",
    re.IGNORECASE,
)

# Generic LAN без бренда СПЕЦЛАН (prod 27.07 SUPR ТЗ): U/UTP cat 5e 2x2x0.52 PE
# U/UTP, F/UTP, S/FTP, SF/TP (без бренда СПЕЦЛАН)
_GENERIC_LAN_SHIELD = r"(?:(?:[USF]/)?/?UTP|S/?FTP|SF/?TP)"
_GENERIC_LAN_SHEATH = r"(?:PE|PVC|LSZH|ZH|PP|FRHF|FRLS)"
_GENERIC_LAN_MARK_PATTERN = re.compile(
    r"("
    rf"{_GENERIC_LAN_SHIELD}"
    r"\s+cat\s*\d\w?"
    r"\s+\d+\s*[xх]\s*\d+(?:\s*[xх]\s*[\d.,]+)?"
    rf"(?:\s+{_GENERIC_LAN_SHEATH})?"
    r")",
    re.IGNORECASE,
)

# Нумерованный список марок в письме (кириллица + латинский OCR)
_LETTER_MARKS_BLOCK = re.compile(
    r"(?:марк[аи]|mapkax?)\s+(?:кабел[ья]|kabena?)[:\s]+"
    r"(.+?)(?=Последующ|Nocnegy|С\s+уважением|Суважением|$)",
    re.IGNORECASE | re.DOTALL,
)
_LETTER_MARK_ITEM = re.compile(
    r"\d+\.\s*"
    rf"({_LAN_BRAND}\s+.+?)"
    r"(?:\s+(?:ТУ|TU|Ty)\s*\d+\.[КкKk]\d{2,3}-\d{3}-\d{4})?"
    r"(?=\s+(?:В\s+количестве|KonuyectBe)|;|\d+\.\s+(?:СПЕЦ|SPECLAN|CMEL)|\d+\.\s+[А-ЯЁA-Z]|$)",
    re.IGNORECASE,
)
_TU_TAIL_PATTERN = re.compile(
    r"ТУ\s*\d+\.[КкKk]\d{2,3}-\d{3}-\d{4}",
    re.IGNORECASE,
)

_MARK_PATTERN = re.compile(
    rf"{_NAME_PART}\s+{_SIZE_PART}",
    re.IGNORECASE,
)

# КГРвЭСТ 3*35+16/3в+3*2,5 - 1140 (направления в ИЛ, звёздочная запись жил)
_STAR_CORE_MARK_PATTERN = re.compile(
    r"(КГ[А-ЯЁа-яёA-Za-z]+\s+"
    r"\d+\*[\d.,]+"
    r"(?:\+[\d\*\/вВa-zA-Z.,]+)*"
    r"(?:\s*-\s*1140)?)",
    re.IGNORECASE,
)

# FLEXICORE (заявка Флексикор, таблица приложения)
# Fire-safety: нг(A) / OCR ur(A)|ng(A) — нормализуется в _fix_series_cable_ocr
_FLEXICORE_FIRE = r"(?:нг|ur|ng)\s*\(\s*[AaАа]\s*\)"
_FLEXICORE_KV = r"(?:кВ|kB|KB|кB|kВ)"
_FLEXICORE_MARK_PATTERN = re.compile(
    r"(FLEXICORE(?:®)?\s+"
    r"(?:"
    rf"FLAT\s+{_FLEXICORE_FIRE}\s*-\s*LS"
    r"|Li(?:YY|YCY)"
    r"|\d{2,}\s+(?:H|CH|CY)"
    rf"(?:\s+{_FLEXICORE_FIRE}\s*-\s*(?:LS|HF))?"
    rf"(?:\s+\d+[.,]?\d*\s*/\s*\d+\s*{_FLEXICORE_KV})?"
    r"|100"
    rf"(?:\s+\d+[.,]?\d*\s*/\s*\d+\s*{_FLEXICORE_KV})?"
    rf"(?:\s+{_FLEXICORE_FIRE}\s*-\s*LS)?"
    rf"(?:\s+\d+[.,]?\d*\s*/\s*\d+\s*{_FLEXICORE_KV})?"
    r"|110"
    rf"(?:\s+{_FLEXICORE_FIRE}\s*-\s*LS)?"
    r"))",
    re.IGNORECASE,
)
_FLEXICORE_TABLE_LINE = re.compile(
    r"(FLEXICORE(?:®)?\s+(?:(?!FLEXICORE|H07RN)[^\n|]){1,70}?)"
    r"(?=\s*\||\s*T[УYYuу]\s|\s*ТУ|\s+FLEXICORE|\s+H07RN|\n|$)",
    re.IGNORECASE,
)
# H07RN-F RU — часто внизу таблицы; OCR: H07RN F RU / H07RNF / HO7RN-F (O≈0)
_H07RN_MARK_PATTERN = re.compile(
    r"((?:H|Н)[0OО]?7(?:RN|RМ|R[NM])[\s\-]*F(?:\s+RU)?)",
    re.IGNORECASE,
)

# VicabFLEX (латиница, направления в ИЛ)
_VICABFLEX_FIRE = r"нг\([AaАа]\)"
_VICABFLEX_MARK_PATTERN = re.compile(
    r"(VicabFLEX\s+"
    r"\d+\s+"
    r"(?:CY|СУ)\s+"
    r"(?:"
    rf"{_VICABFLEX_FIRE}-LS\s+\d+\s*[xх]\s*[\d.,]+(?:\s*\d+/\d+\s*В)?"
    r"|"
    r"\d+[.,]?\d*\s*/\s*\d+\s*кВ\s+\d+\s*[xх]\s*[\d.,]+(?:\s*\d+/\d+\s*В)?"
    r"))",
    re.IGNORECASE,
)

# Имя + размер: пробел опционален (prod: «марки ЛПМФм10х0,08»)
# Только в контексте «марки …» — не в глобальном _MARK_PATTERN (иначе FP).
_NAME_SIZE_CAPTURE = rf"({_NAME_PART}\s*{_SIZE_PART})"

# Контекстный поиск после «марки:» / «марка»
_AFTER_MARKI_PATTERN = re.compile(
    rf"(?:кабел[ья]\s+)?(?:силовой\s+)?марк[аи]\s*:?\s*"
    rf"{_NAME_SIZE_CAPTURE}",
    re.IGNORECASE,
)

# Провод/кабель «марки X» (включая OCR «ПровоА», «Nposoa Mapkn»)
_PRODUCT_MARK_PATTERN = re.compile(
    rf"(?:Прово\w*|Кабел\w*|Mapkn\w*)[^\n]{{0,50}}?марк[аи]\s*:?\s*"
    rf"{_NAME_SIZE_CAPTURE}",
    re.IGNORECASE,
)

_REJECT_PREFIXES = re.compile(
    r"^(?:солнечного|излучения|воздействию|стойкость|требования|наименование|"
    r"марка|образец|№|п\.|пункт|гост|ту|нд)\b",
    re.IGNORECASE,
)

DEFAULT_OCR_DPI = 200
# Практика 09.07: сканы стабильнее на 400 DPI (Tesseract)
SCAN_OCR_DPI = 400
# EasyOCR (PyTorch): тот же класс DPI; opt-in, не default
EASYOCR_OCR_DPI = 400

_EASYOCR_READER = None


def _require_pdfplumber() -> None:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber не установлен. Выполни: pip install pdfplumber")


def _normalize_text(text: str) -> str:
    """Приводит текст PDF/OCR к удобному для поиска виду."""
    text = text.replace("\xa0", " ")
    text = text.replace("—", "-").replace("–", "-")
    text = re.sub(r"-\s*\n\s*", "-", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _normalize_text_for_marks(text: str) -> str:
    """
    Нормализация для поиска марок.

    Сохраняет латинское «x» в размерах (2x2x0,52, Cat 5e) и в F/UTP.
    Кириллическое «х» унифицирует в «x» для единого паттерна.
    """
    text = _normalize_text(text)
    text = text.replace("×", "x").replace("Х", "x").replace("х", "x")
    return text


def _fix_lan_letter_ocr(text: str) -> str:
    """Правки OCR в гарантийных письмах (латиница вместо кириллицы)."""
    text = re.sub(
        r"\b(?:CMELVIAH|CMELAN|SPECLAN|Cneu\w*lan|Sneu\w*lan|СNЕ[UW][АA]{1,2}Н|СNЕLWIАН)\b",
        "СПЕЦЛАН",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"ТараНТ\w+ое\s+ннсбМо", "Гарантийное письмо", text, flags=re.IGNORECASE)
    text = re.sub(r"Мроснм\s+Бас\s+нроБестн", "Просим Вас провести", text, flags=re.IGNORECASE)
    text = re.sub(r"МаркКах\s+Ка6ена", "Марках кабеля", text, flags=re.IGNORECASE)
    text = re.sub(r"мз\s+ТеНеранбНому\s+АннрекТору", "Генеральному директору", text, flags=re.IGNORECASE)
    text = re.sub(r"Испытательный центр", "Испытательный центр", text, flags=re.IGNORECASE)
    text = re.sub(r"МНН/КNN", "ИНН/КПП", text, flags=re.IGNORECASE)
    text = re.sub(
        r"\bMapkax?\s+Kabena?\b",
        "Марках кабеля",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bKonuyectBe\b", "В количестве", text, flags=re.IGNORECASE)
    text = re.sub(
        r"ZH\s+(?:ur|Hr|hr|ng|нr|Нг)\s*\(\s*A\s*\)",
        "ZH нг(А)",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"TapaHTuiHoe\s+nucbmMo",
        "Гарантийное письмо",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"Mpocum\s+Bac\s+nprovectu",
        "Просим Вас провести",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"TeHepanbHomy\s+AnpekTopy",
        "Генеральному директору",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bOOO\s+HNN\b", "ООО НПП", text, flags=re.IGNORECASE)
    # Proizvoditel / Сненка* → см. client_profiles.local.yaml (Спецкабель и др.)
    from .client_profiles import apply_org_ocr_aliases

    text = apply_org_ocr_aliases(text)
    text = re.sub(r"Ispytatelnyj centr", "Испытательный центр", text, flags=re.IGNORECASE)
    text = re.sub(r"\bYa\.\s+", "ул. ", text, flags=re.IGNORECASE)
    return text


def _fix_periodic_letter_ocr(text: str) -> str:
    """Правки OCR в письмах на периодические испытания (таблица марок, периодические)."""
    text = re.sub(
        r"N?Mpocum\s+Bac\s+nprovectm?",
        "Просим Вас провести",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"TeEHEPAAbHOMY\s+ANPeKTOPy", "Генеральному директору", text, flags=re.IGNORECASE)
    text = re.sub(r"KaayxKckni\s+KaGeAbHbIN\s+3GB0A", "Кабельный завод", text, flags=re.IGNORECASE)
    text = re.sub(
        r"Nnepuoauyeckie\s+UCNblITAHMA",
        "периодические испытания",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bBBI-MHr\(A\)", "ВВГнг(А)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bBBI-", "ВВГ-", text, flags=re.IGNORECASE)
    # «ББР-Мнг» / «BBP-Mнг» / «ВВГ-П-Nнг» — частый OCR для ВВГ-Пнг (исх 163)
    text = re.sub(r"ББР-Мнг\(А\)", "ВВГ-Пнг(А)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bББР-Мнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bББР-", "ВВГ-П", text, flags=re.IGNORECASE)
    text = re.sub(r"\bBBP-Mнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bBBP-", "ВВГ-П", text, flags=re.IGNORECASE)
    text = re.sub(r"ВВГ-П-Nнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"ВВГ-П\s*Nнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bВВГ-Мнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bВВГ-Nнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bВВГ\s*Nнг", "ВВГ-Пнг", text, flags=re.IGNORECASE)
    # OCR теряет «П» у плоского: «ВВГнг(А) 3х2,5ок» → «ВВГ-Пнг…»
    text = re.sub(
        r"\bВВГнг(\s*\(\s*[АA]\s*\))\s*(\d+\s*[хx]\s*[\d.,]*ок)",
        r"ВВГ-Пнг\1 \2",
        text,
        flags=re.IGNORECASE,
    )
    # Пробел между fire-safety и размером: «нг(А)3х» → «нг(А) 3х»
    text = re.sub(
        r"([нг]{1,2}\s*\(\s*[АA]\s*\))\s*(\d+\s*[хx])",
        r"\1 \2",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bБбI?NнуСК", "выпускается", text, flags=re.IGNORECASE)
    text = re.sub(r"3[хx]2,50К", "3х2,5ок", text, flags=re.IGNORECASE)
    text = re.sub(r"3[хx]2[,.]50[оoОO]?[кkКK]", "3х2,5ок", text, flags=re.IGNORECASE)
    # Без пробела после запятой — иначе _SIZE_PART обрывает марку на «(N,»
    text = re.sub(r"\(N,\s*РЕ\)", "(N,PE)", text, flags=re.IGNORECASE)
    text = re.sub(r"\(N,\s*PE\)", "(N,PE)", text, flags=re.IGNORECASE)
    text = re.sub(r"--0,66", "-0,66", text)
    text = re.sub(r"NBCur\(A\)-LS", "ПВСнг(А)-LS", text, flags=re.IGNORECASE)
    text = re.sub(r"МБСнг\(А\)\}", "ПВСнг(А)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bМБСнг", "ПВСнг", text, flags=re.IGNORECASE)
    text = re.sub(r"\bAllyB\b", "АПуВ", text, flags=re.IGNORECASE)
    text = re.sub(r"FIБББ", "ПБГВВ", text, flags=re.IGNORECASE)
    text = re.sub(r"FIБIББ", "ПБГВВ", text, flags=re.IGNORECASE)
    text = re.sub(r"FЕРББ", "ПБГВВ", text, flags=re.IGNORECASE)
    text = re.sub(r"\bNBIBB\b", "ПБГВВ", text, flags=re.IGNORECASE)
    text = re.sub(r"F[I1l][БBЕEРP]{2,4}(?=\s*\d|\b)", "ПБГВВ", text, flags=re.IGNORECASE)
    # Пробел «ПБГВВ3х» → «ПБГВВ 3х»
    text = re.sub(r"(ПБГВВ)(\d+\s*[хx])", r"\1 \2", text, flags=re.IGNORECASE)
    text = re.sub(r"Мул\s+Бгат\s*\(А\)", "ПуПВнг(А)", text, flags=re.IGNORECASE)
    text = re.sub(r"Мул\s*Бг?ат\s*\(А\)", "ПуПВнг(А)", text, flags=re.IGNORECASE)
    text = re.sub(r"Пу[ПГ]Внг\s*\(\s*А\s*\)", "ПуПВнг(А)", text, flags=re.IGNORECASE)
    # «…нг(А)-LSLТх 1х6» без бренда → ПуПВнг
    text = re.sub(
        r"(?<![А-ЯЁA-Z])нг\s*\(\s*А\s*\)\s*-\s*LSL[ТT][хxX]\s*(\d+\s*[хx]\s*[\d.,]+)",
        r"ПуПВнг(А)-LSLTx \1",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"LSLТ[хx]", "LSLTx", text)
    text = re.sub(r"LSL[ТT][хxX]", "LSLTx", text)
    text = re.sub(
        r"(?<!ПуПВнг\(А\)-)(?<!-)LSLTx\s*(\d+\s*[хx]\s*[\d.,]+)",
        r"ПуПВнг(А)-LSLTx \1",
        text,
    )
    text = re.sub(r"3[хx]1\.5", "3х1,5", text, flags=re.IGNORECASE)
    text = re.sub(r"3x40K", "3х4ок", text, flags=re.IGNORECASE)
    text = re.sub(r"Nposoa\s+Mapkn\w*", "Провод марки", text, flags=re.IGNORECASE)
    text = re.sub(r"флросоа\s+марКн", "Провод марки", text, flags=re.IGNORECASE)
    text = re.sub(r"Ka6eAb\s+CHACBON\s+MAPK:", "Кабель силовой марки:", text, flags=re.IGNORECASE)
    text = re.sub(r"KaGег?б\s+Сн?АОБОМ\s+МАРКК?:", "Кабель силовой марки:", text, flags=re.IGNORECASE)
    text = re.sub(
        r"Nроснм\s+Бас\s+нросестн",
        "Просим Вас провести",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"нерноануескн\w*", "периодические", text, flags=re.IGNORECASE)
    return text


def _fix_series_cable_ocr(text: str) -> str:
    """Правки OCR в заявках FLEXICORE (в т.ч. после auto-orient стр. 2)."""
    text = re.sub(r"FL[Ee3][XХx][I1l]CORE", "FLEXICORE", text, flags=re.IGNORECASE)
    text = re.sub(r"FLEX1CORE", "FLEXICORE", text, flags=re.IGNORECASE)
    # «H ur(A)» / «Hur(A)» / «CH ur(A)» — склеенный OCR fire-safety
    text = re.sub(
        r"\b(H|CH|CY)(?:\s*)(?:ur|xr|nr|ng|uг|хr)\s*\(\s*A\s*\)",
        r"\1 нг(A)",
        text,
        flags=re.IGNORECASE,
    )
    # Tesseract/EasyOCR: «нг(A)» → ur/ng/nr/нr/xr/nurf после preprocess
    text = re.sub(
        r"\b(?:ur|ng|nr|nurf|uг|xr|хr|xг|нr|нr)\s*\(\s*A\s*\)",
        "нг(A)",
        text,
        flags=re.IGNORECASE,
    )
    # EasyOCR often: «нr(A)» / «нR(A)» / «nг(A)»
    text = re.sub(r"[нn][rгR]\s*\(\s*A\s*\)", "нг(A)", text, flags=re.IGNORECASE)
    text = re.sub(r"нг\s*\(\s*A\s*\)", "нг(A)", text, flags=re.IGNORECASE)
    # EasyOCR: LiYCY / FLAT splits
    text = re.sub(r"\bLi\s*Y\s*CY\b", "LiYCY", text, flags=re.IGNORECASE)
    text = re.sub(r"\bLi\s*YY\b", "LiYY", text, flags=re.IGNORECASE)
    text = re.sub(r"\bFL\s*AT\b", "FLAT", text, flags=re.IGNORECASE)
    text = re.sub(
        r"0[,.]6\s*/\s*1\s*(?:kB|KB|кB|kВ|xB|хB|хВ|xВ)\b",
        "0,6/1 кВ",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\b(?:H|Н)[0OО]?7(?:RN|RМ|R[NM])[\s\-]*F(?:\s*RU)?\b",
        "H07RN-F RU",
        text,
        flags=re.IGNORECASE,
    )
    return text


def _fix_direction_marks_ocr(text: str) -> str:
    """Правки OCR в направлениях в ИЛ (VicabFLEX, CY, кВ)."""
    text = re.sub(
        r"V[иiI][сc]аб[A-Za-zА-Яа-яЁё]{0,10}",
        "VicabFLEX",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"VicabFLEx", "VicabFLEX", text, flags=re.IGNORECASE)
    text = re.sub(
        r"(VicabFLEX\s+\d+)\s+СУ\b",
        r"\1 CY",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"0,6/1\s+КБ\s+", "0,6/1 кВ ", text, flags=re.IGNORECASE)
    text = re.sub(r"600/1000\s+Б\b", "600/1000В", text, flags=re.IGNORECASE)
    text = re.sub(r"300/500\s+Б\b", "300/500 В", text, flags=re.IGNORECASE)
    text = re.sub(r"нг\s*\(\s*A\s*\)", "нг(A)", text, flags=re.IGNORECASE)
    return text


def _fix_latin_brand_ocr(text: str) -> str:
    return _fix_series_cable_ocr(_fix_direction_marks_ocr(text))


def _fix_ocr_confusables(text: str) -> str:
    """Исправляет типичные OCR-ошибки перед поиском марок и ТУ."""
    text = _fix_latin_brand_ocr(text)
    text = _fix_lan_letter_ocr(text)
    text = _fix_periodic_letter_ocr(text)
    return normalize_ocr_text(text)


def is_plausible_mark(mark: str) -> bool:
    """Отсекает явные ложные срабатывания из «рваного» текста PDF/OCR."""
    if re.search(r"\d{4,},\d{2}", mark):
        return False
    if len(re.findall(r"\d{1,3}(?:\s\d{3})*,\d{2}", mark)) >= 2:
        return False
    if re.search(r"\d{2}\.\d{4}", mark):
        return False
    if _REJECT_PREFIXES.search(mark):
        return False
    if re.match(r"^нг\(", mark, re.IGNORECASE):
        return False
    if re.match(r"^кВ\s*\d", mark, re.IGNORECASE):
        return False
    if not re.match(r"^[А-ЯЁA-Z]", mark):
        return False
    if not re.match(
        r"^(?:ККЗ\s+МК\s+|СПЕЦЛАН\s+|VicabFLEX\s+|КГ[А-ЯЁа-яё]|[А-ЯЁA-Z])",
        mark,
        re.IGNORECASE,
    ):
        return False
    has_cyr_size = re.search(
        r"\d+\s*[зЗпП]?\s*[хx]\s*(?:[\d.,]|[а-яёa-zA-Z])", mark, re.IGNORECASE
    )
    has_lan_size = re.search(_SIZE_PART_LATIN, mark, re.IGNORECASE)
    has_star_size = re.search(r"\d+\*[\d.,]", mark)
    is_latin_brand = re.match(r"^(?:FLEXICORE|VicabFLEX|H07RN-F)", mark, re.IGNORECASE)
    is_generic_lan = bool(
        re.match(rf"^{_GENERIC_LAN_SHIELD}\b", mark, re.IGNORECASE)
    )
    if (
        not has_cyr_size
        and not has_lan_size
        and not has_star_size
        and not is_latin_brand
        and not is_generic_lan
    ):
        return False
    # One brand token per mark (joined OCR lines are not a single mark)
    if mark.upper().count("FLEXICORE") > 1:
        return False
    if mark.upper().count("VICABFLEX") > 1:
        return False
    # Generic LAN: «U/UTP cat 5e 2x2…» — name до size, не до «cat 5»
    if is_generic_lan:
        name = re.split(r"\s+\d+\s*[xх]", mark, maxsplit=1, flags=re.IGNORECASE)[0]
    else:
        name = re.split(r"\s+\d", mark, maxsplit=1)[0]
    if len(name) > 100 or len(name) < 3:
        return False
    return True


def _clean_mark(raw: str) -> str:
    """
    Дешёвая детерминированная нормализация кандидата в марку.

    Только strip / regex / локальные OCR-правки. Без SQLite, MarkCorrector,
    fuzzy и Ollama — иначе O(regex-кандидаты × справочник) на каждый finditer.
    Ассистент (fuzzy/LLM) — отдельно, после дедупа, в GUI/CLI.
    """
    mark = raw.strip(" .,;:\n")
    mark = re.sub(
        rf"^(.+?{_SIZE_PART})\s+м\s+\d.*",
        r"\1",
        mark,
        flags=re.IGNORECASE,
    )
    mark = re.sub(r"\s+Упаковка.*$", "", mark, flags=re.IGNORECASE)
    mark = re.sub(
        r"\s+(?:Прово\w*|Кабел\w*|марк[аи]|ТУ|TY|ГОСТ|СТО|Прошу|наименование).*$",
        "",
        mark,
        flags=re.IGNORECASE,
    )
    mark = re.sub(r"\s+\d+\s*м\b.*$", "", mark, flags=re.IGNORECASE)
    if mark.upper().startswith("СПЕЦЛАН"):
        mark = re.sub(r"\s+В\s+количестве.*$", "", mark, flags=re.IGNORECASE)
    if mark.upper().startswith("VICABFLEX"):
        mark = re.sub(r"нг\([Аа]\)", "нг(A)", mark, flags=re.IGNORECASE)
        mark = re.sub(r"\bСУ\b", "CY", mark, flags=re.IGNORECASE)
        mark = re.sub(r"(\d)кВ", r"\1 кВ", mark, flags=re.IGNORECASE)
        return mark.strip()
    if mark.upper().startswith("FLEXICORE") or re.match(r"^(?:H|Н)[0OО]?7", mark, re.I):
        mark = mark.replace("®", "")
        mark = _fix_series_cable_ocr(mark)
        mark = re.sub(r"нг\([Аа]\)", "нг(A)", mark, flags=re.IGNORECASE)
        mark = re.sub(r"\s*_\s*$", "", mark)
        mark = re.sub(r"\s+", " ", mark).strip()
        mark = re.sub(r"\s+H07RN.*$", "", mark, flags=re.IGNORECASE)
        # normalize H07RN variants to canonical GT form
        if re.match(r"^(?:H|Н)[0OО]?7", mark, re.I):
            mark = "H07RN-F RU"
        return mark
    if re.match(r"^КГ[А-ЯЁа-яё]", mark):
        return mark.strip()
    # Локальные OCR-фиксы без БД (латиница→кириллица, fire-class, spacing)
    return normalize_mark_after_ocr(mark.strip(), known_brands=None)


def _context_snippet(text: str, start: int, end: int, radius: int = 120) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    snippet = text[left:right].strip()
    return re.sub(r"\s+", " ", snippet)


def _add_match(
    matches: list[CableMarkMatch],
    seen: set[str],
    mark: str,
    text: str,
    start: int,
    end: int,
    *,
    document: str | None = None,
) -> None:
    mark = _clean_mark(mark)
    key = mark.lower()
    if len(mark) < 5 or key in seen or not is_plausible_mark(mark):
        return
    seen.add(key)
    context = _context_snippet(text, start, end, radius=180)
    doc = document or extract_document_from_text(context)
    matches.append(
        CableMarkMatch(
            mark=mark,
            context=context,
            document=doc,
        )
    )


def _find_letter_list_marks(text: str) -> list[tuple[str, int, int, str | None]]:
    """Марки из нумерованного списка в гарантийном/сопроводительном письме."""
    found: list[tuple[str, int, int, str | None]] = []
    block = _LETTER_MARKS_BLOCK.search(text)
    search_in = block.group(1) if block else text
    base_offset = block.start(1) if block else 0

    for m in _LETTER_MARK_ITEM.finditer(search_in):
        mark = m.group(1).strip()
        if not re.match(r"^(?:СПЕЦЛАН|SPECLAN|CMEL)", mark, re.IGNORECASE):
            continue
        doc = m.group(2).strip() if m.lastindex and m.lastindex >= 2 and m.group(2) else None
        if not doc:
            tail = search_in[m.end() : m.end() + 80]
            tu_m = _TU_TAIL_PATTERN.search(tail)
            doc = tu_m.group(0) if tu_m else None
        if doc:
            doc = extract_document_from_text(doc) or doc
        found.append((mark, base_offset + m.start(1), base_offset + m.end(1), doc))

    if not found:
        for m in _LAN_MARK_PATTERN.finditer(text):
            tail = text[m.end() : m.end() + 80]
            tu_m = _TU_TAIL_PATTERN.search(tail)
            doc = extract_document_from_text(tu_m.group(0)) if tu_m else None
            found.append((m.group(1), m.start(1), m.end(1), doc))

    return found


def _mark_dedupe_key(mark: str) -> str:
    return re.sub(r"\s+", "", mark.lower()).replace("x", "х")


def _dedupe_cable_matches(
    items: list[CableMarkMatch],
    *,
    light: bool = False,
) -> list[CableMarkMatch]:
    """
    Дедупликация марок.

    light=True — только OCR-нормализация (без fuzzy snap по cable_marks).
    Используется для писем на периодические испытания.
    """
    seen: set[str] = set()
    out: list[CableMarkMatch] = []
    for m in items:
        if light:
            fixed = normalize_mark_after_ocr(m.mark.strip())
        else:
            fixed = _clean_mark(m.mark)
        key = _mark_dedupe_key(fixed)
        if fixed and key not in seen:
            seen.add(key)
            out.append(
                CableMarkMatch(
                    mark=fixed,
                    context=m.context,
                    document=m.document,
                    requirements_raw=m.requirements_raw,
                )
            )
    return out


_FLEXICORE_FIRE_TAIL = re.compile(
    r"^(FLEXICORE(?:®)?\s+.+?)\s+нг\(A\)-(LS|HF)$",
    re.IGNORECASE,
)
_FLEXICORE_VOLT_TAIL = re.compile(
    r"^(FLEXICORE(?:®)?\s+.+?)\s+(\d+[.,]?\d*\s*/\s*\d+\s*кВ)$",
    re.IGNORECASE,
)


def _expand_flexicore_combinations(items: list[CableMarkMatch]) -> list[CableMarkMatch]:
    """Если OCR дал «100 нг(A)-LS» и «100 0,6/1 кВ» порознь — добавляем полную строку.

    Не brand-hardcode: только комбинирование уже найденных FLEXICORE-кандидатов.
    """
    by_key = {m.mark.lower(): m for m in items}
    fires: dict[str, tuple[str, str, CableMarkMatch]] = {}
    volts: dict[str, tuple[str, str, CableMarkMatch]] = {}
    for m in items:
        if not m.mark.upper().startswith("FLEXICORE"):
            continue
        fm = _FLEXICORE_FIRE_TAIL.match(m.mark)
        if fm:
            base = re.sub(r"\s+", " ", fm.group(1)).strip()
            fires[base.lower()] = (base, fm.group(2).upper().replace("А", "A"), m)
            continue
        vm = _FLEXICORE_VOLT_TAIL.match(m.mark)
        if vm:
            base = re.sub(r"\s+", " ", vm.group(1)).strip()
            volt = vm.group(2).replace(".", ",")
            volts[base.lower()] = (base, volt, m)

    extras: list[CableMarkMatch] = []
    for key, (base, fire_suf, src) in fires.items():
        if key not in volts:
            continue
        _, volt, _vsrc = volts[key]
        combined = f"{base} нг(A)-{fire_suf} {volt}"
        combined = _clean_mark(combined)
        ckey = combined.lower()
        if ckey in by_key or not is_plausible_mark(combined):
            continue
        by_key[ckey] = CableMarkMatch(
            mark=combined,
            context=src.context,
            document=src.document,
            requirements_raw=src.requirements_raw,
        )
        extras.append(by_key[ckey])
    return items + extras


def _extract_marks_from_family(text: str, family_id: str) -> list[CableMarkMatch]:
    """Извлекает марки через YAML-семейство, только при уверенном совпадении."""
    from .families.registry import DocumentFamily, get_family_registry

    registry = get_family_registry()
    family: DocumentFamily | None = registry.get(family_id)
    if family is None or not family.is_confident_match(text):
        return []

    if family_id == "periodic_letter_v1":
        from .periodic_letter_extractor import extract_marks_from_periodic_letter

        periodic = extract_marks_from_periodic_letter(text)
        cleaned = _dedupe_cable_matches(periodic, light=True)
        if len(cleaned) >= family.min_marks_threshold:
            return cleaned
        return []

    matches: list[CableMarkMatch] = []
    seen: set[str] = set()
    for _kind, pattern in family.mark_patterns:
        for m in pattern.finditer(text):
            raw = m.group(1) if m.lastindex else m.group(0)
            _add_match(matches, seen, raw, text, m.start(), m.end())

    if family_id == "lan_letter_v1" and len(matches) < family.min_marks_threshold:
        for mark, start, end, doc in _find_letter_list_marks(text):
            _add_match(matches, seen, mark, text, start, end, document=doc)

    if len(matches) >= family.min_marks_threshold:
        return matches
    return []


def find_cable_marks(text: str) -> list[CableMarkMatch]:
    """
    Ищет марки кабелей по структурному паттерну «название + NхM».

    Не использует список брендов — любая строка нужной формы.
    Семейства YAML (периодические, LAN) — только при score ≥ confidence_threshold.
    """
    base = _normalize_text_for_marks(text)
    if not base:
        return []

    seen: set[str] = set()
    matches: list[CableMarkMatch] = []

    # Латиница и КГ* — до normalize_ocr_text (он портит FLEXICORE/VicabFLEX)
    latin_safe = _fix_latin_brand_ocr(base)
    for pattern in (
        _FLEXICORE_MARK_PATTERN,
        _FLEXICORE_TABLE_LINE,
        _H07RN_MARK_PATTERN,
        _VICABFLEX_MARK_PATTERN,
        _STAR_CORE_MARK_PATTERN,
    ):
        for m in pattern.finditer(latin_safe):
            raw = m.group(1) if m.lastindex else m.group(0)
            _add_match(matches, seen, raw, latin_safe, m.start(), m.end())

    normalized = _fix_ocr_confusables(base)

    from .families.registry import get_family_registry

    family = get_family_registry().detect_best(normalized)
    if family is not None:
        family_marks = _extract_marks_from_family(normalized, family.id)
        if family_marks:
            if family.id == "periodic_letter_v1":
                return _expand_flexicore_combinations(family_marks)
            return _expand_flexicore_combinations(
                _dedupe_cable_matches(matches + family_marks)
            )

    for mark, start, end, doc in _find_letter_list_marks(normalized):
        _add_match(matches, seen, mark, normalized, start, end, document=doc)

    for pattern in (
        _LAN_MARK_PATTERN,
        _GENERIC_LAN_MARK_PATTERN,
        _AFTER_MARKI_PATTERN,
        _PRODUCT_MARK_PATTERN,
        _MARK_PATTERN,
    ):
        for m in pattern.finditer(normalized):
            raw = m.group(1) if m.lastindex else m.group(0)
            _add_match(matches, seen, raw, normalized, m.start(), m.end())

    return _expand_flexicore_combinations(matches)


OcrEngineChoice = Literal["auto", "tesseract", "easyocr"]
OcrEngineResolved = Literal["tesseract", "easyocr", "none"]


def _easyocr_importable() -> bool:
    try:
        import easyocr  # noqa: F401

        return True
    except ImportError:
        return False


def resolve_ocr_engine(
    preferred: OcrEngineChoice | str = "auto",
) -> OcrEngineResolved:
    """Выбор OCR-движка.

    * ``auto`` — tesseract, иначе easyocr (если установлен)
    * ``tesseract`` / ``easyocr`` — предпочтение; при отсутствии — fallback
      на другой доступный движок (чтобы на prod не «долбить» мёртвый easyocr)
    """
    pref = (preferred or "auto").strip().lower()
    has_tess = bool(_find_tesseract())
    has_easy = _easyocr_importable()

    if pref in ("easyocr", "pytorch_cv", "pytorch", "torch"):
        if has_easy:
            return "easyocr"
        if has_tess:
            logger.warning(
                "EasyOCR/torch недоступен — fallback на Tesseract",
            )
            return "tesseract"
        return "none"
    if pref == "tesseract":
        if has_tess:
            return "tesseract"
        if has_easy:
            logger.warning("Tesseract не найден — fallback на EasyOCR")
            return "easyocr"
        return "none"
    # auto: Tesseract first (лёгкий, default на prod)
    if has_tess:
        return "tesseract"
    if has_easy:
        return "easyocr"
    return "none"


def _pdf_cache_fingerprint(path: Path) -> str:
    stat = path.stat()
    payload = f"{path.resolve()}|{stat.st_size}|{stat.st_mtime_ns}"
    return hashlib.sha256(payload.encode()).hexdigest()[:24]


def _ocr_cache_path(path: Path, dpi: int, engine: str, *, preprocess: str | None = None) -> Path:
    fingerprint = _pdf_cache_fingerprint(path)
    safe_stem = re.sub(r"[^\w\-.]+", "_", path.stem)[:48]
    pre_tag = f"_pre{preprocess}" if preprocess else ""
    return OCR_CACHE_DIR / f"{safe_stem}_{fingerprint}_dpi{dpi}_{engine}{pre_tag}.txt"


def _read_ocr_cache(path: Path, dpi: int, engine: str, *, preprocess: str | None = None) -> str | None:
    cache_path = _ocr_cache_path(path, dpi, engine, preprocess=preprocess)
    if not cache_path.is_file():
        return None
    text = cache_path.read_text(encoding="utf-8")
    return text if text.strip() else None


def _write_ocr_cache(
    path: Path, dpi: int, engine: str, text: str, *, preprocess: str | None = None
) -> None:
    if not text.strip():
        return
    cache_path = _ocr_cache_path(path, dpi, engine, preprocess=preprocess)
    cache_path.parent.mkdir(parents=True, exist_ok=True)
    cache_path.write_text(text, encoding="utf-8")


def clear_ocr_cache() -> int:
    """Удаляет файлы кэша OCR. Возвращает число удалённых файлов."""
    if not OCR_CACHE_DIR.is_dir():
        return 0
    removed = 0
    for cache_file in OCR_CACHE_DIR.glob("*.txt"):
        cache_file.unlink(missing_ok=True)
        removed += 1
    return removed


def _get_easyocr_reader():
    global _EASYOCR_READER
    if _EASYOCR_READER is None:
        import easyocr

        _EASYOCR_READER = easyocr.Reader(["ru", "en"], gpu=False, verbose=False)
    return _EASYOCR_READER


def _find_tesseract() -> str | None:
    """Ищет tesseract.exe: env → PATH → portable проекта → Program Files.

    На рабочем ПК путь часто отличается от машины разработчика.
    Задайте явно (любой из вариантов):

    * ``TESSERACT_CMD`` / ``TESSERACT_PATH`` — полный путь к ``tesseract.exe``
    * положите portable: ``tools/Tesseract-OCR/tesseract.exe``
    * добавьте папку Tesseract в системный PATH
    """
    import os

    for env_key in ("TESSERACT_CMD", "TESSERACT_PATH"):
        raw = (os.environ.get(env_key) or "").strip().strip('"')
        if raw:
            p = Path(raw)
            if p.is_file():
                return str(p)
            # иногда указывают папку, а не exe
            if p.is_dir():
                for name in ("tesseract.exe", "tesseract"):
                    cand = p / name
                    if cand.is_file():
                        return str(cand)

    found = shutil.which("tesseract")
    if found:
        return found
    project_root = Path(__file__).resolve().parents[2]
    for candidate in (
        project_root / "tools" / "Tesseract-OCR" / "tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"D:\Tesseract-OCR\tesseract.exe",
        r"D:\Apps\Tesseract-OCR\tesseract.exe",
        r"D:\Program Files\Tesseract-OCR\tesseract.exe",
    ):
        path = Path(candidate)
        if path.exists():
            return str(path)
    return None


def _render_pages(pdf_path: Path, dpi: int = DEFAULT_OCR_DPI) -> list:
    """Рендерит страницы PDF в PIL.Image через PyMuPDF."""
    import fitz
    from PIL import Image

    doc = fitz.open(str(pdf_path))
    images: list = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    for page in doc:
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        images.append(Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples))
    doc.close()
    return images


def _resolve_preprocess_tag(use_preprocess: bool) -> str | None:
    if not use_preprocess:
        return None
    from .ocr.preprocess import PREPROCESS_VERSION, is_cv_available

    return PREPROCESS_VERSION if is_cv_available() else None


def _ocr_bottom_band_eng(image, pytesseract, *, band_ratio: float = 0.22) -> str:
    """OCR нижней полосы страницы eng-only — латиница H07RN / FLEXICORE footer.

    Геометрия (crop), не brand-regex: Tesseract rus+eng часто «съедает» хвост таблицы.
    """
    from PIL import Image

    if image is None:
        return ""
    w, h = image.size
    y0 = max(0, int(h * (1.0 - band_ratio)))
    crop = image.crop((0, y0, w, h))
    # slight upscale helps thin Latin glyphs
    if crop.height < 200:
        crop = crop.resize((crop.width * 2, crop.height * 2), Image.Resampling.LANCZOS)
    config = "--psm 6 --oem 1"
    try:
        return pytesseract.image_to_string(crop, lang="eng", config=config) or ""
    except Exception:
        return ""


def _ocr_image_tesseract(image, pytesseract, *, use_preprocess: bool = True) -> str:
    # rus первичен — кириллица в марках; eng для Cat/UTP в LAN-кабеле
    # Auto-orient (OSD) first — FLEXICORE стр. 2 скана повёрнута на 90°.
    # Adaptive threshold after orient often *hurts* Latin brand tables → soft preprocess
    # + merge with raw so RU letters and FLEXICORE both survive.
    # Bottom-band eng pass (35s): recovers H07RN-F RU when full-page OCR misses footer.
    from .ocr.preprocess import correct_orientation, is_cv_available, preprocess_for_ocr

    tesseract_cmd = getattr(pytesseract.pytesseract, "tesseract_cmd", None)
    image = correct_orientation(image, tesseract_cmd=tesseract_cmd or None)
    config = "--psm 6 --oem 1"
    text_raw = pytesseract.image_to_string(image, lang="rus+eng", config=config)
    text_footer = _ocr_bottom_band_eng(image, pytesseract)
    parts: list[str] = []
    if not use_preprocess or not is_cv_available():
        if text_raw.strip():
            parts.append(text_raw.strip())
        if text_footer.strip() and text_footer.strip() not in (text_raw or ""):
            parts.append(text_footer.strip())
        return "\n\n".join(parts)
    # Soft path: deskew/denoise/upscale without adaptive threshold
    soft = preprocess_for_ocr(image, adaptive_threshold=False)
    text_soft = pytesseract.image_to_string(soft, lang="rus+eng", config=config)
    if not text_soft.strip():
        primary = text_raw
    elif not text_raw.strip():
        primary = text_soft
    elif len(text_soft) > len(text_raw) * 1.15:
        primary = f"{text_soft}\n\n{text_raw}".strip() if text_raw.strip() else text_soft
    else:
        primary = f"{text_raw}\n\n{text_soft}".strip() if text_soft.strip() != text_raw.strip() else text_raw
    if text_footer.strip() and text_footer.strip() not in (primary or ""):
        primary = f"{primary}\n\n{text_footer}".strip()
    return primary


def _ocr_with_tesseract(
    pdf_path: Path,
    dpi: int = DEFAULT_OCR_DPI,
    *,
    use_preprocess: bool = True,
    progress=None,
) -> str:
    import pytesseract

    from .progress import NULL_PROGRESS

    prog = progress or NULL_PROGRESS
    tesseract_cmd = _find_tesseract()
    if not tesseract_cmd:
        raise RuntimeError(
            "Tesseract OCR не найден. Установи: "
            "https://github.com/UB-Mannheim/tesseract/wiki "
            "(нужен язык rus)."
        )
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    prog("Рендер страниц PDF…", stage="ocr")
    images = _render_pages(pdf_path, dpi=dpi)
    n = len(images)
    if n == 1:
        prog("OCR страница 1/1…", current=0, total=1, stage="ocr")
        text = _ocr_image_tesseract(images[0], pytesseract, use_preprocess=use_preprocess)
        prog("OCR страница 1/1 готово", current=1, total=1, stage="ocr")
        return text.strip()

    # С progress — постранично (видно остаток); без — параллельно быстрее
    if progress is not None:
        parts: list[str] = []
        for idx, img in enumerate(images):
            prog(
                f"OCR страница {idx + 1}/{n}…",
                current=idx,
                total=n,
                stage="ocr",
            )
            parts.append(
                _ocr_image_tesseract(img, pytesseract, use_preprocess=use_preprocess).strip()
            )
            prog(
                f"OCR страница {idx + 1}/{n} готово",
                current=idx + 1,
                total=n,
                stage="ocr",
            )
        return "\n\n".join(p for p in parts if p)

    parts_p: list[str] = [""] * n
    with ThreadPoolExecutor(max_workers=min(4, n)) as pool:
        futures = {
            pool.submit(_ocr_image_tesseract, img, pytesseract, use_preprocess=use_preprocess): idx
            for idx, img in enumerate(images)
        }
        for future in as_completed(futures):
            idx = futures[future]
            parts_p[idx] = future.result().strip()

    return "\n\n".join(p for p in parts_p if p)


def _easyocr_prepare_image(image):
    """Подготовка кадра для EasyOCR: orient + контраст/upscale (без жёсткого threshold)."""
    import numpy as np

    from .ocr.preprocess import correct_orientation, is_cv_available

    tesseract_cmd = _find_tesseract()
    image = correct_orientation(image, tesseract_cmd=tesseract_cmd)
    if not is_cv_available():
        return np.array(image.convert("RGB"))

    import cv2

    arr = np.array(image.convert("RGB"))
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    # CLAHE — выравнивание контраста (сканы с серой таблицей)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)
    h, w = gray.shape[:2]
    # EasyOCR любит крупные глифы; до ~2200 px по высоте
    if h < 2200:
        scale = 2200 / h
        gray = cv2.resize(
            gray,
            (int(w * scale), int(h * scale)),
            interpolation=cv2.INTER_CUBIC,
        )
    # лёгкое усиление краёв букв
    blur = cv2.GaussianBlur(gray, (0, 0), 1.0)
    sharp = cv2.addWeighted(gray, 1.4, blur, -0.4, 0)
    return cv2.cvtColor(sharp, cv2.COLOR_GRAY2RGB)


def _easyocr_lines_from_image(reader, rgb_arr) -> list[str]:
    """Построчный rec (не paragraph) — таблицы лучше, чем «слипшийся» абзац."""
    # detail=1 → (bbox, text, conf); paragraph=False — строки таблицы не склеиваются
    raw = reader.readtext(
        rgb_arr,
        detail=1,
        paragraph=False,
        mag_ratio=1.5,
        text_threshold=0.6,
        low_text=0.3,
        link_threshold=0.3,
        canvas_size=2560,
        contrast_ths=0.1,
        adjust_contrast=0.5,
    )
    if not raw:
        return []
    # sort by vertical position then x
    items: list[tuple[float, float, str]] = []
    for bbox, text, conf in raw:
        if not text or not str(text).strip():
            continue
        if conf is not None and float(conf) < 0.15:
            continue
        ys = [p[1] for p in bbox]
        xs = [p[0] for p in bbox]
        items.append((sum(ys) / len(ys), min(xs), str(text).strip()))
    items.sort(key=lambda t: (round(t[0] / 12.0), t[1]))  # group ~12px rows
    # merge tokens on same row band
    lines: list[str] = []
    band: list[tuple[float, str]] = []
    band_y: float | None = None
    for y, x, text in items:
        if band_y is None or abs(y - band_y) <= 18:
            band.append((x, text))
            band_y = y if band_y is None else (band_y * 0.7 + y * 0.3)
        else:
            band.sort(key=lambda t: t[0])
            lines.append(" ".join(t for _, t in band))
            band = [(x, text)]
            band_y = y
    if band:
        band.sort(key=lambda t: t[0])
        lines.append(" ".join(t for _, t in band))
    return lines


def _ocr_with_easyocr(
    pdf_path: Path, dpi: int = EASYOCR_OCR_DPI, *, progress=None
) -> str:
    """EasyOCR (PyTorch CV): orient + CLAHE/upscale + line-mode + footer band."""
    from .progress import NULL_PROGRESS

    prog = progress or NULL_PROGRESS
    prog("Загрузка EasyOCR (PyTorch)…", stage="ocr")
    reader = _get_easyocr_reader()
    images = _render_pages(pdf_path, dpi=dpi)
    n = len(images)
    parts: list[str] = []
    for idx, image in enumerate(images):
        prog(
            f"torch-CV страница {idx + 1}/{n}…",
            current=idx,
            total=n,
            stage="ocr",
        )
        rgb = _easyocr_prepare_image(image)
        lines = _easyocr_lines_from_image(reader, rgb)
        h = rgb.shape[0]
        y0 = int(h * 0.78)
        footer = rgb[y0:, :, :]
        foot_lines = _easyocr_lines_from_image(reader, footer)
        page_text = "\n".join(lines)
        if foot_lines:
            foot_text = "\n".join(foot_lines)
            if foot_text.strip() and foot_text.strip() not in page_text:
                page_text = f"{page_text}\n{foot_text}".strip()
        if page_text.strip():
            parts.append(_fix_series_cable_ocr(page_text))
        prog(
            f"torch-CV страница {idx + 1}/{n} готово",
            current=idx + 1,
            total=n,
            stage="ocr",
        )
    return "\n\n".join(parts)


def ocr_pdf(
    pdf_path: Path | str,
    dpi: int = DEFAULT_OCR_DPI,
    *,
    use_cache: bool = True,
    use_preprocess: bool = True,
    engine: OcrEngineChoice | str = "auto",
) -> str:
    """
    Распознаёт текст со скана.

    ``engine``: auto | tesseract | easyocr (pytorch_cv).
    Кэш: ``data/ocr_cache/`` — отдельный ключ на engine (+ preprocess tag).
    """
    text, _used = ocr_pdf_ex(
        pdf_path,
        dpi=dpi,
        use_cache=use_cache,
        use_preprocess=use_preprocess,
        engine=engine,
    )
    return text


def ocr_pdf_ex(
    pdf_path: Path | str,
    dpi: int = DEFAULT_OCR_DPI,
    *,
    use_cache: bool = True,
    use_preprocess: bool = True,
    engine: OcrEngineChoice | str = "auto",
    progress=None,
) -> tuple[str, OcrEngineResolved]:
    """Как ``ocr_pdf``, но возвращает (text, фактический engine)."""
    from .progress import NULL_PROGRESS

    prog = progress or NULL_PROGRESS
    path = Path(pdf_path)
    preferred = (engine or "auto").strip().lower()
    if preferred in ("pytorch_cv", "pytorch", "torch"):
        preferred = "easyocr"
    if preferred not in ("auto", "tesseract", "easyocr"):
        preferred = "auto"

    resolved = resolve_ocr_engine(preferred)  # type: ignore[arg-type]
    if resolved == "none":
        raise RuntimeError(
            "OCR недоступен.\n"
            "1) Установите Tesseract OCR + language pack rus "
            "(или задайте TESSERACT_CMD=…\\tesseract.exe).\n"
            "2) Опционально: pip install easyocr (тяжёлый, не default).\n"
            "3) Для Word-заявок OCR не нужен — откройте .docx."
        )

    preprocess_tag = (
        _resolve_preprocess_tag(use_preprocess) if resolved == "tesseract" else None
    )

    if use_cache:
        cached = _read_ocr_cache(path, dpi, resolved, preprocess=preprocess_tag)
        if cached is not None:
            logger.debug(
                "OCR cache hit: %s (%s, dpi=%d, pre=%s)",
                path.name,
                resolved,
                dpi,
                preprocess_tag,
            )
            prog("OCR из кэша (мгновенно)", current=1, total=1, stage="ocr")
            return normalize_ocr_text(cached), resolved

    text = ""
    used: OcrEngineResolved = resolved

    if resolved == "tesseract":
        try:
            text = _ocr_with_tesseract(
                path, dpi=dpi, use_preprocess=use_preprocess, progress=progress
            )
        except Exception as exc:
            logger.warning("Tesseract OCR failed: %s", exc)
            if preferred != "auto":
                raise
            used = "easyocr"
            preprocess_tag = None
            if use_cache:
                cached = _read_ocr_cache(path, dpi, "easyocr")
                if cached is not None:
                    return cached, "easyocr"

    need_easy = resolved == "easyocr" or (
        preferred == "auto" and used == "easyocr" and not text.strip()
    )
    if need_easy or (preferred == "auto" and resolved == "tesseract" and not text.strip()):
        try:
            text = _ocr_with_easyocr(path, dpi=dpi, progress=progress)
            used = "easyocr"
            preprocess_tag = None
        except ImportError as exc:
            if preferred == "easyocr":
                raise RuntimeError(
                    "Для EasyOCR (PyTorch CV): pip install easyocr (venv на D:)"
                ) from exc
            logger.warning("EasyOCR fallback unavailable: %s", exc)

    if text.strip():
        text = normalize_ocr_text(text)

    if use_cache and text.strip():
        _write_ocr_cache(
            path,
            dpi,
            used,
            text,
            preprocess=preprocess_tag if used == "tesseract" else None,
        )
    return text, used


def extract_text(pdf_path: Path | str, *, use_ocr: bool = False) -> str:
    """Извлекает текст: pdfplumber, для сканов — с OCR если use_ocr=True."""
    _require_pdfplumber()
    path = Path(pdf_path)
    parts: list[str] = []

    with pdfplumber.open(path) as doc:
        for page in doc.pages:
            page_text = page.extract_text() or ""
            if page_text:
                parts.append(page_text)

    text = "\n\n".join(parts)
    if not text.strip() and use_ocr:
        text = ocr_pdf(path)
    return text


def extract_tables(pdf_path: Path | str) -> list[list[list[str]]]:
    """Извлекает таблицы. Каждая ячейка — строка (пустые → '')."""
    _require_pdfplumber()
    path = Path(pdf_path)
    tables: list[list[list[str]]] = []

    with pdfplumber.open(path) as doc:
        for page in doc.pages:
            for table in page.extract_tables() or []:
                cleaned = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in table
                ]
                if any(any(cell for cell in row) for row in cleaned):
                    tables.append(cleaned)

    return tables


def _tables_to_text(tables: list[list[list[str]]]) -> str:
    return "\n".join(" ".join(cell for cell in row if cell) for table in tables for row in table)


def _resolve_cable_marks(
    text: str,
    tables: list[list[list[str]]],
) -> list[CableMarkMatch]:
    """
    Марки кабелей: table-first для направлений, иначе regex по тексту.

    Таблица направления приоритетнее плоского текста PDF/Word.
    Текст перед regex сжимается (Word-шаблоны с merge иначе → секунды CPU).
    """
    if tables:
        from .direction_table_extractor import extract_marks_from_tables

        table_marks = extract_marks_from_tables(tables)
        if table_marks:
            return table_marks

    search_text = build_search_text(text, tables)
    return find_cable_marks(_compact_text_for_marks(search_text))


def _cyrillic_letter_ratio(text: str) -> float:
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return 0.0
    cyrillic = sum(1 for c in letters if "\u0400" <= c <= "\u04FF")
    return cyrillic / len(letters)


def _pdf_page_stats(pdf_path: Path) -> tuple[int, bool, str]:
    """page_count, has_images, joined pdfplumber text."""
    _require_pdfplumber()
    with pdfplumber.open(pdf_path) as doc:
        page_count = len(doc.pages)
        has_images = any(page.images for page in doc.pages)
        parts = [page.extract_text() or "" for page in doc.pages]
        return page_count, has_images, "\n\n".join(parts).strip()


def _needs_ocr_fallback(text: str, *, has_images: bool) -> bool:
    """
    Текстовый слой есть, но нечитаем (латиница вместо кириллицы) — нужен OCR.

    Типично для PDF со сканом-вложением и мусорным pseudo-text layer.
    """
    if not has_images or not text.strip():
        return False
    return _cyrillic_letter_ratio(text) < 0.2


def _detect_scanned(pdf_path: Path) -> tuple[bool, int]:
    """True, если PDF похож на скан (есть изображения, но нет текста)."""
    page_count, has_images, text = _pdf_page_stats(pdf_path)
    is_scanned = page_count > 0 and not text and has_images
    return is_scanned, page_count


def build_search_text(text: str, tables: list[list[list[str]]] | None = None) -> str:
    """Текст заявки + плоское представление таблиц для поиска org/марок."""
    if not tables:
        return _dedupe_consecutive_lines(text)
    joined = f"{text}\n{_tables_to_text(tables)}".strip()
    return _dedupe_consecutive_lines(joined)


def extract_from_pdf(
    pdf_path: Path | str,
    *,
    use_ocr: bool = True,
    ocr_dpi: int = DEFAULT_OCR_DPI,
    use_ocr_cache: bool = True,
    ocr_engine: OcrEngineChoice | str = "auto",
    progress=None,
) -> PdfExtractionResult:
    """
    Полное извлечение: текст, таблицы, марки кабелей, метаданные.

    Для сканов при use_ocr=True автоматически запускает OCR.
    ``ocr_engine``: auto | tesseract | easyocr (pytorch_cv) — A/B spike 35v.
    ``progress``: callback(message, current=, total=, stage=) для GUI.
    """
    from .progress import NULL_PROGRESS

    prog = progress or NULL_PROGRESS
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF не найден: {path}")

    prog("Анализ PDF…", current=0, total=100, stage="open")
    is_scanned, page_count = _detect_scanned(path)
    _page_count, has_images, plumber_text = _pdf_page_stats(path)
    ocr_used = False
    ocr_engine_used: str | None = None
    tables: list[list[list[str]]] = []

    table_ocr_text = ""
    use_ocr_path = is_scanned or (
        use_ocr and has_images and _needs_ocr_fallback(plumber_text, has_images=has_images)
    )

    if use_ocr_path:
        eng_pref = (ocr_engine or "auto").strip().lower()
        if eng_pref in ("easyocr", "pytorch_cv", "pytorch", "torch"):
            scan_dpi = max(ocr_dpi, EASYOCR_OCR_DPI)
        elif is_scanned:
            scan_dpi = max(ocr_dpi, SCAN_OCR_DPI)
        else:
            scan_dpi = ocr_dpi
        prog(
            f"OCR ({eng_pref}, DPI {scan_dpi}, ~{page_count} стр.)…",
            current=5,
            total=100,
            stage="ocr",
        )
        text, ocr_engine_used = ocr_pdf_ex(
            path,
            dpi=scan_dpi,
            use_cache=use_ocr_cache,
            use_preprocess=use_ocr,
            engine=ocr_engine,
            progress=progress,
        )
        ocr_used = True
        if not is_scanned:
            logger.info(
                "PDF %s: pseudo-text layer (cyr=%.0f%%) — OCR fallback",
                path.name,
                _cyrillic_letter_ratio(plumber_text) * 100,
            )
        if not use_ocr:
            logger.warning("PDF %s: скан — OCR включён принудительно", path.name)
        if is_scanned:
            try:
                from .ocr.table import TABLE_OCR_DPI, ocr_tables_from_pdf, tables_text_from_results

                prog("OCR таблиц…", current=70, total=100, stage="tables")
                table_dpi = max(scan_dpi, TABLE_OCR_DPI)
                table_results = ocr_tables_from_pdf(path, dpi=table_dpi)
                table_ocr_text = tables_text_from_results(table_results)
                if table_ocr_text:
                    tables = [result.rows for result in table_results if result.rows]
                    text = f"{text}\n\n{table_ocr_text}".strip()
                    logger.info(
                        "PDF %s: table OCR added %d chars from %d page(s)",
                        path.name,
                        len(table_ocr_text),
                        len(table_results),
                    )
            except Exception as exc:
                logger.debug("Table OCR skipped for %s: %s", path.name, exc)
    else:
        prog("Чтение текстового слоя…", current=20, total=100, stage="text")
        text = extract_text(path)
        tables = extract_tables(path)

    prog("Поиск марок…", current=85, total=100, stage="marks")
    search_text = build_search_text(text, tables)
    cable_marks = _resolve_cable_marks(text, tables)
    prog("Организации…", current=92, total=100, stage="orgs")
    organizations = finalize_organizations(extract_organizations(search_text), search_text)
    prog("Готово", current=100, total=100, stage="done")

    logger.info(
        "PDF %s: pages=%d, text=%d chars, tables=%d, marks=%d, orgs=%d, scanned=%s, ocr=%s, engine=%s",
        path.name,
        page_count,
        len(text),
        len(tables),
        len(cable_marks),
        len(organizations),
        is_scanned,
        ocr_used,
        ocr_engine_used,
    )

    return PdfExtractionResult(
        source_path=str(path.resolve()),
        source_type="pdf",
        page_count=page_count,
        text=text,
        tables=tables,
        cable_marks=cable_marks,
        organizations=organizations,
        customer_name=pick_customer_name(organizations),
        manufacturer_name=pick_manufacturer_name(organizations),
        is_scanned=is_scanned,
        ocr_used=ocr_used,
        ocr_engine=ocr_engine_used,
        extracted_at=datetime.now(),
    )


def _collapse_horizontal_merge_cells(cells: list[str]) -> list[str]:
    """Убрать дубли ячеек от горизонтального merge (python-docx повторяет text)."""
    if not cells:
        return []
    out: list[str] = []
    prev: str | None = None
    for cell in cells:
        if prev is not None and cell == prev:
            continue
        out.append(cell)
        prev = cell
    return out


def _dedupe_consecutive_lines(text: str) -> str:
    """Схлопнуть подряд идущие одинаковые строки (шаблоны Word / merge)."""
    if not text:
        return ""
    out: list[str] = []
    prev: str | None = None
    blank_run = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            blank_run += 1
            if blank_run <= 1:
                out.append("")
            prev = ""
            continue
        blank_run = 0
        if stripped == prev:
            continue
        out.append(stripped)
        prev = stripped
    return "\n".join(out).strip()


def _unique_lines_preserve_order(text: str, *, max_chars: int = 60_000) -> str:
    """Уникальные непустые строки (порядок сохраняется) — для regex по маркам."""
    if not text:
        return ""
    seen: set[str] = set()
    lines: list[str] = []
    size = 0
    for line in text.splitlines():
        key = line.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        if size + len(key) + 1 > max_chars:
            break
        lines.append(key)
        size += len(key) + 1
    return "\n".join(lines)


def _compact_text_for_marks(text: str, *, max_chars: int = 60_000) -> str:
    """Сжать текст перед find_cable_marks (шаблоны Word иначе → 60k+ и секунды regex)."""
    compact = _dedupe_consecutive_lines(text)
    if len(compact) <= max_chars:
        return compact
    unique = _unique_lines_preserve_order(compact, max_chars=max_chars)
    return unique if unique else compact[:max_chars]


def load_docx_content(docx_path: Path | str) -> tuple[str, list[list[list[str]]]]:
    """
    Один проход по .docx: абзацы + таблицы.

    Горизонтальные merge-ячейки схлопываются (иначе python-docx отдаёт
    один и тот же текст 7–11 раз → 60k+ символов и долгий regex).
    """
    from docx import Document

    path = Path(docx_path)
    doc = Document(str(path))

    para_parts: list[str] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t:
            para_parts.append(t)
    text = _dedupe_consecutive_lines("\n".join(para_parts))

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        cleaned: list[list[str]] = []
        prev_row_key: tuple[str, ...] | None = None
        for row in table.rows:
            raw = [cell.text.strip() for cell in row.cells]
            collapsed = _collapse_horizontal_merge_cells(raw)
            if not any(collapsed):
                continue
            # подряд одинаковые строки (вертикальный merge / повтор шаблона)
            row_key = tuple(collapsed)
            if row_key == prev_row_key:
                continue
            prev_row_key = row_key
            cleaned.append(collapsed)
        if cleaned:
            tables.append(cleaned)
    return text, tables


def extract_text_from_docx(docx_path: Path | str) -> str:
    """Извлекает текст из Word (.docx): абзацы + схлопнутые таблицы."""
    text, tables = load_docx_content(docx_path)
    return build_search_text(text, tables)


def extract_tables_from_docx(docx_path: Path | str) -> list[list[list[str]]]:
    """Таблицы Word (merge-ячейки схлопнуты) — для table-first направлений."""
    _text, tables = load_docx_content(docx_path)
    return tables


def _build_extraction_result(
    *,
    source_path: Path,
    source_type: Literal["pdf", "docx", "text"],
    text: str,
    page_count: int = 0,
    tables: list[list[list[str]]] | None = None,
    is_scanned: bool = False,
    ocr_used: bool = False,
) -> PdfExtractionResult:
    """Собирает результат: марки + организации из текста заявки (без ассистента)."""
    t0 = time.perf_counter()
    # Word-направления: абзацы часто пустые (весь текст в таблицах).
    # result.text должен содержать таблицы — иначе validate/type/HITL видят ".".
    full_text = build_search_text(text, tables) if tables else (text or "")
    search_text = _compact_text_for_marks(full_text)
    t_compact = time.perf_counter()
    cable_marks = _resolve_cable_marks(text, tables or [])
    t_marks = time.perf_counter()
    organizations = finalize_organizations(
        extract_organizations(search_text), search_text
    )
    t_orgs = time.perf_counter()
    resolved = source_path if source_path.is_absolute() or source_path.exists() else source_path
    try:
        path_str = str(resolved.resolve())
    except OSError:
        path_str = str(source_path)
    logger.info(
        "parse timing file=%s marks=%.3fs orgs=%.3fs compact=%.3fs total=%.3fs "
        "n_marks=%d n_orgs=%d search_chars=%d result_text_chars=%d para_chars=%d tables=%d",
        Path(path_str).name,
        t_marks - t_compact,
        t_orgs - t_marks,
        t_compact - t0,
        t_orgs - t0,
        len(cable_marks),
        len(organizations),
        len(search_text),
        len(full_text),
        len(text or ""),
        len(tables or []),
    )
    return PdfExtractionResult(
        source_path=path_str,
        source_type=source_type,
        page_count=page_count,
        text=full_text,
        tables=tables or [],
        cable_marks=cable_marks,
        organizations=organizations,
        customer_name=pick_customer_name(organizations),
        manufacturer_name=pick_manufacturer_name(organizations),
        is_scanned=is_scanned,
        ocr_used=ocr_used,
        extracted_at=datetime.now(),
    )


def extract_from_text(
    text: str,
    *,
    source_label: str = "free_text",
) -> PdfExtractionResult:
    """
    Вход «речь/текст заказчика» или любой свободный текст заявки.

    Без OCR: сразу извлекает марки, организации и вид испытаний-эвристики.
    """
    cleaned = (text or "").strip()
    if not cleaned:
        raise ValueError("Пустой текст заявки")
    # Виртуальный source_path — не файл на диске
    label = re.sub(r"[^\w.\-]+", "_", source_label, flags=re.UNICODE).strip("_") or "free_text"
    result = _build_extraction_result(
        source_path=Path(f"{label}.txt"),
        source_type="text",
        text=cleaned,
        page_count=1,
    )
    return result.model_copy(update={"source_path": f"text://{label}"})


def extract_from_document(
    path: Path | str,
    *,
    use_ocr: bool = True,
    ocr_dpi: int = DEFAULT_OCR_DPI,
    use_ocr_cache: bool = True,
    ocr_engine: OcrEngineChoice | str = "auto",
    progress=None,
) -> PdfExtractionResult:
    """
    Единая точка входа для заявок: PDF или Word (.docx).

    PDF делегирует в extract_from_pdf; Word — текст + марки + организации.
    ``ocr_engine``: auto | tesseract | easyocr (pytorch_cv).
    """
    from .progress import NULL_PROGRESS

    prog = progress or NULL_PROGRESS
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_from_pdf(
            file_path,
            use_ocr=use_ocr,
            ocr_dpi=ocr_dpi,
            use_ocr_cache=use_ocr_cache,
            ocr_engine=ocr_engine,
            progress=progress,
        )
    if suffix == ".docx":
        # Word: OCR не нужен. Один проход Document (не два), merge-ячейки схлопнуты.
        # Ассистент/fuzzy/Ollama сюда не входят — только детерминированный разбор.
        t_start = time.perf_counter()
        prog("Чтение Word…", current=10, total=100, stage="text")
        text, tables = load_docx_content(file_path)
        t_open = time.perf_counter()
        prog(
            f"Поиск марок… (текст {len(text)} симв., таблиц {len(tables)})",
            current=70,
            total=100,
            stage="marks",
        )
        result = _build_extraction_result(
            source_path=file_path,
            source_type="docx",
            text=text,
            tables=tables,
            page_count=1,
        )
        t_done = time.perf_counter()
        prog("Готово", current=100, total=100, stage="done")
        logger.info(
            "docx extract file=%s open=%.3fs build=%.3fs total=%.3fs "
            "para_chars=%s result_text_chars=%s tables=%s marks=%s customer_len=%s",
            file_path.name,
            t_open - t_start,
            t_done - t_open,
            t_done - t_start,
            len(text),
            len(result.text or ""),
            len(tables),
            len(result.cable_marks),
            len(result.customer_name or ""),
        )
        return result
    raise ValueError(f"Неподдерживаемый формат: {suffix}. Используйте PDF или .docx")